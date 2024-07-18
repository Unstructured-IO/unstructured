import datetime
import os
import pathlib

import docx
import openpyxl
import pytest

import unstructured.file_utils.metadata as meta
from test_unstructured.unit_utils import example_doc_path

DIRECTORY = pathlib.Path(__file__).parent.resolve()
EXAMPLE_JPG_FILENAME = example_doc_path("img/example.jpg")


def test_get_docx_metadata_from_filename(tmpdir):
    filename = os.path.join(tmpdir, "test-doc.docx")

    document = docx.Document()
    document.add_paragraph("Lorem ipsum dolor sit amet.")
    document.core_properties.author = "Mr. Miagi"
    document.save(filename)

    metadata = meta.get_docx_metadata(filename=filename)
    assert metadata.author == "Mr. Miagi"
    assert metadata.to_dict()["author"] == "Mr. Miagi"


def test_get_docx_metadata_from_file(tmpdir):
    filename = os.path.join(tmpdir, "test-doc.docx")

    document = docx.Document()
    document.add_paragraph("Lorem ipsum dolor sit amet.")
    document.core_properties.author = "Mr. Miagi"
    document.save(filename)

    with open(filename, "rb") as f:
        metadata = meta.get_docx_metadata(file=f)
    assert metadata.author == "Mr. Miagi"


def test_get_docx_metadata_raises_without_file_or_filename():
    with pytest.raises(FileNotFoundError):
        meta.get_docx_metadata()


def test_get_xlsx_metadata_from_filename(tmpdir):
    filename = os.path.join(tmpdir, "test-excel.xlsx")

    workbook = openpyxl.Workbook()
    workbook.properties.creator = "Mr. Miagi"
    workbook.save(filename)

    metadata = meta.get_xlsx_metadata(filename=filename)
    metadata.author = "Mr. Miagi"


def test_get_xlsx_metadata_from_file(tmpdir):
    filename = os.path.join(tmpdir, "test-excel.xlsx")

    workbook = openpyxl.Workbook()
    workbook.properties.creator = "Mr. Miagi"
    workbook.save(filename)

    with open(filename, "rb") as f:
        metadata = meta.get_xlsx_metadata(file=f)
    metadata.author = "Mr. Miagi"


def test_get_xlsx_metadata_raises_without_file_or_filename():
    with pytest.raises(FileNotFoundError):
        meta.get_xlsx_metadata()


def test_get_jpg_metadata_from_filename():
    metadata = meta.get_jpg_metadata(filename=EXAMPLE_JPG_FILENAME)
    assert metadata.modified == datetime.datetime(2003, 12, 14, 12, 1, 44)
    assert metadata.exif_data["Make"] == "Canon"


def test_get_jpg_metadata_from_file():
    with open(EXAMPLE_JPG_FILENAME, "rb") as f:
        metadata = meta.get_jpg_metadata(file=f)
    assert metadata.modified == datetime.datetime(2003, 12, 14, 12, 1, 44)
    assert metadata.exif_data["Make"] == "Canon"


def test_get_jpg_metadata_raises_without_file_or_filename():
    with pytest.raises(FileNotFoundError):
        meta.get_jpg_metadata()


def test_get_exif_datetime():
    exif_data = {"DateTime": "2022:12:23 15:49:00", "DateTimeOriginal": "2020:12:14 12:00:00"}
    date = meta._get_exif_datetime(exif_data, "DateTime")
    assert date == datetime.datetime(2022, 12, 23, 15, 49, 0)


def test_get_exif_datetime_ignores_bad_formats():
    exif_data = {"DateTime": "2022-12-23TZ15:49:00", "DateTimeOriginal": "2020:12:14 12:00:00"}
    date = meta._get_exif_datetime(exif_data, "DateTime")
    assert date is None


def test_get_exif_datetime_ignores_missing_key():
    exif_data = {"Datetime": "2022-12-23TZ15:49:00", "DateTimeOriginal": "2020:12:14 12:00:00"}
    date = meta._get_exif_datetime(exif_data, "DateTimeDigitized")
    assert date is None
