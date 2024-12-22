# pyright: reportPrivateUsage=false

"""Test suite for `unstructured.partition.docx` module."""

from __future__ import annotations

import hashlib
import io
import pathlib
import re
import tempfile
from typing import Any, Iterator

import docx
import pytest
from docx.document import Document
from docx.text.paragraph import Paragraph
from pytest_mock import MockFixture

from test_unstructured.unit_utils import (
    FixtureRequest,
    Mock,
    assert_round_trips_through_JSON,
    example_doc_path,
    function_mock,
    instance_mock,
    property_mock,
)
from unstructured.chunking.title import chunk_by_title
from unstructured.documents.elements import (
    Address,
    CompositeElement,
    Element,
    Footer,
    Header,
    Image,
    ListItem,
    NarrativeText,
    PageBreak,
    Table,
    TableChunk,
    Text,
    Title,
)
from unstructured.partition.docx import (
    DocxPartitionerOptions,
    _DocxPartitioner,
    partition_docx,
    register_picture_partitioner,
)
from unstructured.partition.utils.constants import (
    UNSTRUCTURED_INCLUDE_DEBUG_METADATA,
    PartitionStrategy,
)

# -- docx-file loading behaviors -----------------------------------------------------------------


def test_partition_docx_from_filename(
    mock_document_file_path: str, expected_elements: list[Element]
):
    elements = partition_docx(mock_document_file_path)

    assert elements == expected_elements
    assert elements[0].metadata.page_number is None
    for element in elements:
        assert element.metadata.filename == "mock_document.docx"
    if UNSTRUCTURED_INCLUDE_DEBUG_METADATA:
        assert {element.metadata.detection_origin for element in elements} == {"docx"}


def test_partition_docx_with_spooled_file(
    mock_document_file_path: str, expected_elements: list[Text]
):
    """`partition_docx()` accepts a SpooledTemporaryFile as its `file` argument.

    `python-docx` will NOT accept a `SpooledTemporaryFile` in Python versions before 3.11 so we need
    to ensure the source file is appropriately converted in this case.
    """
    with tempfile.SpooledTemporaryFile() as spooled_temp_file:
        with open(mock_document_file_path, "rb") as test_file:
            spooled_temp_file.write(test_file.read())
        spooled_temp_file.seek(0)

        elements = partition_docx(file=spooled_temp_file)

    assert elements == expected_elements
    assert all(e.metadata.filename is None for e in elements)


def test_partition_docx_from_file(mock_document_file_path: str, expected_elements: list[Text]):
    with open(mock_document_file_path, "rb") as f:
        elements = partition_docx(file=f)
    assert elements == expected_elements
    for element in elements:
        assert element.metadata.filename is None


def test_partition_docx_uses_file_path_when_both_are_specified(
    mock_document_file_path: str, expected_elements: list[Text]
):
    f = io.BytesIO(b"abcde")
    elements = partition_docx(filename=mock_document_file_path, file=f)
    assert elements == expected_elements


def test_partition_docx_raises_with_neither():
    with pytest.raises(ValueError, match="either `filename` or `file` argument must be provided"):
        partition_docx()


# ------------------------------------------------------------------------------------------------


def test_parition_docx_from_team_chat():
    """Docx with no sections partitions recognizing both paragraphs and tables."""
    elements = partition_docx(example_doc_path("teams_chat.docx"))
    assert [e.text for e in elements] == [
        "0:0:0.0 --> 0:0:1.510\nSome Body\nOK. Yeah.",
        "0:0:3.270 --> 0:0:4.250\nJames Bond\nUmm.",
        "saved-by Dennis Forsythe",
    ]
    assert [type(e) for e in elements] == [Text, Text, Table]


@pytest.mark.parametrize("infer_table_structure", [True, False])
def test_partition_docx_infer_table_structure(infer_table_structure: bool):
    elements = partition_docx(
        example_doc_path("fake_table.docx"), infer_table_structure=infer_table_structure
    )
    table_element_has_text_as_html_field = (
        hasattr(elements[0].metadata, "text_as_html")
        and elements[0].metadata.text_as_html is not None
    )
    assert table_element_has_text_as_html_field == infer_table_structure


def test_partition_docx_processes_table():
    elements = partition_docx(example_doc_path("fake_table.docx"))

    assert isinstance(elements[0], Table)
    assert elements[0].text == ("Header Col 1 Header Col 2 Lorem ipsum A Link example")
    assert elements[0].metadata.text_as_html == (
        "<table>"
        "<tr><td>Header Col 1</td><td>Header Col 2</td></tr>"
        "<tr><td>Lorem ipsum</td><td>A Link example</td></tr>"
        "</table>"
    )
    assert elements[0].metadata.filename == "fake_table.docx"


def test_partition_docx_grabs_header_and_footer():
    elements = partition_docx(example_doc_path("handbook-1p.docx"))

    assert elements[0] == Header("US Trustee Handbook")
    assert elements[-1] == Footer("Copyright")
    for element in elements:
        assert element.metadata.filename == "handbook-1p.docx"


# -- page-break behaviors ------------------------------------------------------------------------


def test_partition_docx_includes_neither_page_breaks_nor_numbers_when_rendered_breaks_not_present():
    """Hard page-breaks by themselves are not enough to locate page-breaks in a document.

    In particular, they are redundant when rendered page-breaks are present, which they usually are
    in a native Word document, so lead to double-counting those page-breaks. When rendered page
    breaks are *not* present, only a small fraction will be represented by hard page-breaks so hard
    breaks are a false-positive and will generally produce incorrect page numbers.
    """
    elements = partition_docx(
        example_doc_path("handbook-1p-no-rendered-page-breaks.docx"), include_page_breaks=True
    )

    assert "PageBreak" not in [type(e).__name__ for e in elements]
    assert all(e.metadata.page_number is None for e in elements)


def test_partition_docx_includes_page_numbers_when_page_break_elements_are_suppressed():
    """Page-number metadata is not supressed when `include_page_breaks` arga is False.

    Only inclusion of PageBreak elements is affected by that option.
    """
    elements = partition_docx(example_doc_path("handbook-1p.docx"), include_page_breaks=False)

    assert "PageBreak" not in [type(e).__name__ for e in elements]
    assert elements[1].metadata.page_number == 1
    assert elements[-2].metadata.page_number == 2


def test_partition_docx_includes_page_break_elements_when_so_instructed():
    elements = partition_docx(
        example_doc_path("handbook-1p.docx"), include_page_breaks=True, starting_page_number=3
    )

    assert "PageBreak" in [type(e).__name__ for e in elements]
    assert elements[1].metadata.page_number == 3
    assert elements[-2].metadata.page_number == 4


# ------------------------------------------------------------------------------------------------


def test_partition_docx_detects_lists():
    elements = partition_docx(example_doc_path("example-list-items-multiple.docx"))

    assert elements[-1] == ListItem(
        "This is simply dummy text of the printing and typesetting industry.",
    )
    assert sum(1 for e in elements if isinstance(e, ListItem)) == 10


# -- .metadata.filename --------------------------------------------------------------------------


def test_partition_docx_from_filename_prefers_metadata_filename_when_provided():
    elements = partition_docx(example_doc_path("simple.docx"), metadata_filename="test")
    assert all(element.metadata.filename == "test" for element in elements)


def test_partition_docx_from_file_prefers_metadata_filename_when_provided():
    with open(example_doc_path("simple.docx"), "rb") as f:
        elements = partition_docx(file=f, metadata_filename="test")
    assert all(element.metadata.filename == "test" for element in elements)


# -- .metadata.last_modified ---------------------------------------------------------------------


def test_partition_docx_from_file_path_gets_last_modified_from_filesystem(mocker: MockFixture):
    filesystem_last_modified = "2029-07-05T09:24:28"
    mocker.patch(
        "unstructured.partition.docx.get_last_modified_date", return_value=filesystem_last_modified
    )

    elements = partition_docx(example_doc_path("fake.docx"))

    assert elements[0].metadata.last_modified == filesystem_last_modified


def test_partition_docx_from_file_gets_last_modified_None():
    with open(example_doc_path("simple.docx"), "rb") as f:
        elements = partition_docx(file=f)

    assert elements[0].metadata.last_modified is None


def test_partition_docx_from_file_path_prefers_metadata_last_modified(mocker: MockFixture):
    filesystem_last_modified = "2023-11-01T14:13:07"
    metadata_last_modified = "2020-07-05T09:24:28"
    mocker.patch(
        "unstructured.partition.docx.get_last_modified_date", return_value=filesystem_last_modified
    )

    elements = partition_docx(
        example_doc_path("fake.docx"), metadata_last_modified=metadata_last_modified
    )

    assert elements[0].metadata.last_modified == metadata_last_modified


def test_partition_docx_from_file_prefers_metadata_last_modified():
    metadata_last_modified = "2020-07-05T09:24:28"
    with open(example_doc_path("simple.docx"), "rb") as f:
        elements = partition_docx(file=f, metadata_last_modified=metadata_last_modified)

    assert elements[0].metadata.last_modified == metadata_last_modified


# ------------------------------------------------------------------------------------------------


def test_get_emphasized_texts_from_paragraph(
    opts_args: dict[str, Any], expected_emphasized_texts: list[dict[str, str]]
):
    opts_args["file_path"] = example_doc_path("fake-doc-emphasized-text.docx")
    opts = DocxPartitionerOptions(**opts_args)
    partitioner = _DocxPartitioner(opts)

    paragraph = partitioner._document.paragraphs[1]
    emphasized_texts = list(partitioner._iter_paragraph_emphasis(paragraph))
    assert paragraph.text == "I am a bold italic bold-italic text."
    assert emphasized_texts == expected_emphasized_texts

    paragraph = partitioner._document.paragraphs[2]
    emphasized_texts = list(partitioner._iter_paragraph_emphasis(paragraph))
    assert paragraph.text == ""
    assert emphasized_texts == []

    paragraph = partitioner._document.paragraphs[3]
    emphasized_texts = list(partitioner._iter_paragraph_emphasis(paragraph))
    assert paragraph.text == "I am a normal text."
    assert emphasized_texts == []


def test_iter_table_emphasis(
    opts_args: dict[str, Any], expected_emphasized_texts: list[dict[str, str]]
):
    opts_args["file_path"] = example_doc_path("fake-doc-emphasized-text.docx")
    opts = DocxPartitionerOptions(**opts_args)
    partitioner = _DocxPartitioner(opts)
    table = partitioner._document.tables[0]

    emphasized_texts = list(partitioner._iter_table_emphasis(table))

    assert emphasized_texts == expected_emphasized_texts


def test_table_emphasis(
    opts_args: dict[str, Any],
    expected_emphasized_text_contents: list[str],
    expected_emphasized_text_tags: list[str],
):
    opts_args["file_path"] = example_doc_path("fake-doc-emphasized-text.docx")
    opts = DocxPartitionerOptions(**opts_args)
    partitioner = _DocxPartitioner(opts)
    table = partitioner._document.tables[0]

    emphasized_text_contents, emphasized_text_tags = partitioner._table_emphasis(table)

    assert emphasized_text_contents == expected_emphasized_text_contents
    assert emphasized_text_tags == expected_emphasized_text_tags


def test_partition_docx_grabs_emphasized_texts(
    expected_emphasized_text_contents: list[str],
    expected_emphasized_text_tags: list[str],
):
    elements = partition_docx(example_doc_path("fake-doc-emphasized-text.docx"))

    assert isinstance(elements[0], Table)
    assert elements[0].metadata.emphasized_text_contents == expected_emphasized_text_contents
    assert elements[0].metadata.emphasized_text_tags == expected_emphasized_text_tags

    assert elements[1] == NarrativeText("I am a bold italic bold-italic text.")
    assert elements[1].metadata.emphasized_text_contents == expected_emphasized_text_contents
    assert elements[1].metadata.emphasized_text_tags == expected_emphasized_text_tags

    assert elements[2] == NarrativeText("I am a normal text.")
    assert elements[2].metadata.emphasized_text_contents is None
    assert elements[2].metadata.emphasized_text_tags is None


def test_partition_docx_with_json(mock_document_file_path: str):
    elements = partition_docx(mock_document_file_path)
    assert_round_trips_through_JSON(elements)


def test_parse_category_depth_by_style(opts_args: dict[str, Any]):
    opts_args["file_path"] = example_doc_path("category-level.docx")
    opts = DocxPartitionerOptions(**opts_args)
    partitioner = _DocxPartitioner(opts)

    # Category depths are 0-indexed and relative to the category type
    # Title, list item, bullet, narrative text, etc.
    test_cases = [
        (0, "Call me Ishmael."),
        (0, "A Heading 1"),
        (0, "Whenever I find myself growing grim"),
        (0, "A top level list item"),
        (1, "Next level"),
        (1, "Same"),
        (0, "Second top-level list item"),
        (0, "whenever I find myself involuntarily"),
        (0, ""),  # Empty paragraph
        (1, "A Heading 2"),
        (0, "This is my substitute for pistol and ball"),
        (0, "Another Heading 1"),
        (0, "There now is your insular city"),
    ]

    paragraphs = partitioner._document.paragraphs
    for idx, (depth, text) in enumerate(test_cases):
        paragraph = paragraphs[idx]
        actual_depth = partitioner._parse_category_depth_by_style(paragraph)
        assert text in paragraph.text, f"paragraph[{[idx]}].text does not contain {text}"
        assert (
            actual_depth == depth
        ), f"expected paragraph[{idx}] to have depth=={depth}, got {actual_depth}"


def test_parse_category_depth_by_style_name(opts_args: dict[str, Any]):
    opts = DocxPartitionerOptions(**opts_args)
    partitioner = _DocxPartitioner(opts)
    test_cases = [
        (0, "Heading 1"),
        (1, "Heading 2"),
        (2, "Heading 3"),
        (1, "Subtitle"),
        (0, "List"),
        (1, "List 2"),
        (2, "List 3"),
        (0, "List Bullet"),
        (1, "List Bullet 2"),
        (2, "List Bullet 3"),
        (0, "List Number"),
        (1, "List Number 2"),
        (2, "List Number 3"),
    ]

    for idx, (depth, text) in enumerate(test_cases):
        assert (
            partitioner._parse_category_depth_by_style_name(text) == depth
        ), f"test case {test_cases[idx]} failed"


def test_parse_category_depth_by_style_ilvl(opts_args: dict[str, Any]):
    opts = DocxPartitionerOptions(**opts_args)
    partitioner = _DocxPartitioner(opts)
    assert partitioner._parse_category_depth_by_style_ilvl() == 0


def test_add_chunking_strategy_on_partition_docx_default_args():
    chunk_elements = partition_docx(
        example_doc_path("handbook-1p.docx"), chunking_strategy="by_title"
    )
    elements = partition_docx(example_doc_path("handbook-1p.docx"))
    chunks = chunk_by_title(elements)

    assert chunk_elements != elements
    assert chunk_elements == chunks


def test_add_chunking_strategy_on_partition_docx():
    docx_path = example_doc_path("fake-doc-emphasized-text.docx")

    chunk_elements = partition_docx(
        docx_path, chunking_strategy="by_title", max_characters=9, combine_text_under_n_chars=5
    )
    elements = partition_docx(docx_path)
    chunks = chunk_by_title(elements, max_characters=9, combine_text_under_n_chars=5)

    assert chunk_elements == chunks
    assert elements != chunk_elements
    for chunk in chunks:
        assert isinstance(chunk, (CompositeElement, TableChunk))
        assert len(chunk.text) <= 9


# -- language behaviors --------------------------------------------------------------------------


def test_partition_docx_element_metadata_has_languages():
    filename = example_doc_path("handbook-1p.docx")
    elements = partition_docx(filename=filename)
    assert elements[0].metadata.languages == ["eng"]


def test_partition_docx_respects_detect_language_per_element():
    filename = example_doc_path("language-docs/eng_spa_mult.docx")
    elements = partition_docx(filename=filename, detect_language_per_element=True)
    langs = [element.metadata.languages for element in elements]
    assert langs == [["eng"], ["spa", "eng"], ["eng"], ["eng"], ["spa"]]


def test_partition_docx_respects_languages_arg():
    filename = example_doc_path("handbook-1p.docx")
    elements = partition_docx(filename=filename, languages=["deu"])
    assert elements[0].metadata.languages == ["deu"]


def test_partition_docx_raises_TypeError_for_invalid_languages():
    with pytest.raises(TypeError):
        filename = example_doc_path("handbook-1p.docx")
        partition_docx(filename=filename, languages="eng")


# ------------------------------------------------------------------------------------------------


def test_partition_docx_includes_hyperlink_metadata():
    elements = partition_docx(example_doc_path("hlink-meta.docx"))

    # -- regular paragraph, no hyperlinks --
    element = elements[0]
    assert element.text == "One"
    metadata = element.metadata
    assert metadata.links is None
    assert metadata.link_texts is None
    assert metadata.link_urls is None

    # -- paragraph with "internal-jump" hyperlinks, no URL --
    element = elements[1]
    assert element.text == "Two with link to bookmark."
    metadata = element.metadata
    assert metadata.links is None
    assert metadata.link_texts is None
    assert metadata.link_urls is None

    # -- paragraph with external link, no fragment --
    element = elements[2]
    assert element.text == "Three with link to foo.com."
    metadata = element.metadata
    assert metadata.links == [
        {
            "start_index": 11,
            "text": "link to foo.com",
            "url": "https://foo.com",
        },
    ]
    assert metadata.link_texts == ["link to foo.com"]
    assert metadata.link_urls == ["https://foo.com"]

    # -- paragraph with external link that has query string --
    element = elements[3]
    assert element.text == "Four with link to foo.com searching for bar."
    metadata = element.metadata
    assert metadata.links == [
        {
            "start_index": 10,
            "text": "link to foo.com searching for bar",
            "url": "https://foo.com?q=bar",
        },
    ]
    assert metadata.link_texts == ["link to foo.com searching for bar"]
    assert metadata.link_urls == ["https://foo.com?q=bar"]

    # -- paragraph with external link with separate URI fragment --
    element = elements[4]
    assert element.text == "Five with link to foo.com introduction section."
    metadata = element.metadata
    assert metadata.links == [
        {
            "start_index": 10,
            "text": "link to foo.com introduction section",
            "url": "http://foo.com/#intro",
        },
    ]
    assert metadata.link_texts == ["link to foo.com introduction section"]
    assert metadata.link_urls == ["http://foo.com/#intro"]

    # -- paragraph with link to file on local filesystem --
    element = elements[7]
    assert element.text == "Eight with link to file."
    metadata = element.metadata
    assert metadata.links == [
        {
            "start_index": 11,
            "text": "link to file",
            "url": "court-exif.jpg",
        },
    ]
    assert metadata.link_texts == ["link to file"]
    assert metadata.link_urls == ["court-exif.jpg"]

    # -- regular paragraph, no hyperlinks, ensure no state is retained --
    element = elements[8]
    assert element.text == "Nine."
    metadata = element.metadata
    assert metadata.links is None
    assert metadata.link_texts is None
    assert metadata.link_urls is None


def test_partition_docx_assigns_deterministic_and_unique_element_ids():
    document_path = example_doc_path("duplicate-paragraphs.docx")

    ids = [element.id for element in partition_docx(document_path)]
    ids_2 = [element.id for element in partition_docx(document_path)]

    # -- ids match even though partitioned separately (deterministic on content) --
    assert ids == ids_2
    # -- ids are unique --
    assert len(ids) == len(set(ids))


# -- shape behaviors -----------------------------------------------------------------------------


def test_it_considers_text_inside_shapes():
    # -- <bracketed> text is written inside inline shapes --
    partitioned_doc = partition_docx(example_doc_path("docx-shapes.docx"))
    assert [element.text for element in partitioned_doc] == [
        "Paragraph with single <inline-image> within.",
        "Paragraph with <inline-image1> and <inline-image2> within.",
        # -- text "<floating-shape>" in floating shape is ignored --
        "Paragraph with floating shape attached.",
    ]


# -- image sub-partitioning behaviors ------------------------------------------------------------


def test_partition_docx_generates_no_Image_elements_by_default():
    assert not any(
        isinstance(e, Image) for e in partition_docx(example_doc_path("contains-pictures.docx"))
    )


def test_partition_docx_uses_registered_picture_partitioner():
    class FakeParagraphPicturePartitioner:
        @classmethod
        def iter_elements(
            cls, paragraph: Paragraph, opts: DocxPartitionerOptions
        ) -> Iterator[Image]:
            call_hash = hashlib.sha1(f"{paragraph.text}{opts.strategy}".encode()).hexdigest()
            yield Image(f"Image with hash {call_hash}, strategy: {opts.strategy}")

    register_picture_partitioner(FakeParagraphPicturePartitioner)

    elements = partition_docx(example_doc_path("contains-pictures.docx"))

    # -- picture-partitioner registration has module-lifetime, so need to de-register this fake
    # -- so other tests in same test-run don't use it
    DocxPartitionerOptions._PicturePartitionerCls = None

    assert len(elements) == 11
    image_elements = [e for e in elements if isinstance(e, Image)]
    assert len(image_elements) == 6
    assert [e.text for e in image_elements] == [
        "Image with hash 429de54e71f1f0fb395b6f6191961a3ea1b64dc0, strategy: hi_res",
        "Image with hash 5e0cd2c62809377d8ce7422d8ca6b0cf5f4453bc, strategy: hi_res",
        "Image with hash 429de54e71f1f0fb395b6f6191961a3ea1b64dc0, strategy: hi_res",
        "Image with hash ccbd34be6096544babc391890cb0849c24cc046c, strategy: hi_res",
        "Image with hash a41b819c7b4a9750ec0f9198c59c2057d39c653c, strategy: hi_res",
        "Image with hash ba0dc2a1205af8f6d9e06c8d415df096b0a9c428, strategy: hi_res",
    ]


# -- module-level fixtures -----------------------------------------------------------------------


@pytest.fixture()
def expected_elements() -> list[Text]:
    return [
        Title("These are a few of my favorite things:"),
        ListItem("Parrots"),
        ListItem("Hockey"),
        Text("Analysis"),
        NarrativeText("This is my first thought. This is my second thought."),
        NarrativeText("This is my third thought."),
        Text("2023"),
        Address("DOYLESTOWN, PA 18901"),
    ]


@pytest.fixture()
def expected_emphasized_text_contents() -> list[str]:
    return ["bold", "italic", "bold-italic", "bold-italic"]


@pytest.fixture()
def expected_emphasized_text_tags() -> list[str]:
    return ["b", "i", "b", "i"]


@pytest.fixture()
def expected_emphasized_texts():
    return [
        {"text": "bold", "tag": "b"},
        {"text": "italic", "tag": "i"},
        {"text": "bold-italic", "tag": "b"},
        {"text": "bold-italic", "tag": "i"},
    ]


@pytest.fixture()
def mock_document():
    document = docx.Document()

    document.add_paragraph("These are a few of my favorite things:", style="Heading 1")
    # NOTE(robinson) - this should get picked up as a list item due to the •
    document.add_paragraph("• Parrots", style="Normal")
    # NOTE(robinson) - this should get dropped because it's empty
    document.add_paragraph("• ", style="Normal")
    document.add_paragraph("Hockey", style="List Bullet")
    # NOTE(robinson) - this should get dropped because it's empty
    document.add_paragraph("", style="List Bullet")
    # NOTE(robinson) - this should get picked up as a title
    document.add_paragraph("Analysis", style="Normal")
    # NOTE(robinson) - this should get dropped because it is empty
    document.add_paragraph("", style="Normal")
    # NOTE(robinson) - this should get picked up as a narrative text
    document.add_paragraph("This is my first thought. This is my second thought.", style="Normal")
    document.add_paragraph("This is my third thought.", style="Body Text")
    # NOTE(robinson) - this should just be regular text
    document.add_paragraph("2023")
    # NOTE(robinson) - this should be an address
    document.add_paragraph("DOYLESTOWN, PA 18901")

    return document


@pytest.fixture()
def mock_document_file_path(mock_document: Document, tmp_path: pathlib.Path) -> str:
    filename = str(tmp_path / "mock_document.docx")
    mock_document.save(filename)
    return filename


@pytest.fixture()
def opts_args() -> dict[str, Any]:
    """All default arguments for `DocxPartitionerOptions`.

    Individual argument values can be changed to suit each test. Makes construction of opts more
    compact for testing purposes.
    """
    return {
        "file": None,
        "file_path": None,
        "include_page_breaks": True,
        "infer_table_structure": True,
        "strategy": None,
    }


# ================================================================================================
# ISOLATED UNIT TESTS
# ================================================================================================
# These test components used by `partition_docx()` in isolation such that all edge cases can be
# exercised.
# ================================================================================================


class DescribeDocxPartitionerOptions:
    """Unit-test suite for `unstructured.partition.docx.DocxPartitionerOptions` objects."""

    # -- .load() ---------------------------------

    def it_provides_a_validating_constructor(self, opts_args: dict[str, Any]):
        opts_args["file_path"] = example_doc_path("simple.docx")

        opts = DocxPartitionerOptions.load(**opts_args)

        assert isinstance(opts, DocxPartitionerOptions)

    def and_it_raises_when_options_are_not_valid(self, opts_args: dict[str, Any]):
        with pytest.raises(ValueError, match="no DOCX document specified, "):
            DocxPartitionerOptions.load(**opts_args)

    # -- .document -------------------------------

    def it_loads_the_docx_document(
        self,
        request: FixtureRequest,
        opts_args: dict[str, Any],
    ):
        document_ = instance_mock(request, Document)
        docx_Document_ = function_mock(
            request, "unstructured.partition.docx.docx.Document", return_value=document_
        )
        _docx_file_prop_ = property_mock(
            request, DocxPartitionerOptions, "_docx_file", return_value="abcde.docx"
        )
        opts = DocxPartitionerOptions(**opts_args)

        document = opts.document

        _docx_file_prop_.assert_called_once_with()
        docx_Document_.assert_called_once_with("abcde.docx")
        assert document is document_

    # -- .include_page_breaks --------------------

    @pytest.mark.parametrize("arg_value", [True, False])
    def it_knows_whether_to_emit_PageBreak_elements_as_part_of_the_output_element_stream(
        self, arg_value: bool, opts_args: dict[str, Any]
    ):
        opts_args["include_page_breaks"] = arg_value
        opts = DocxPartitionerOptions(**opts_args)

        assert opts.include_page_breaks is arg_value

    # -- .infer_table_structure ------------------

    @pytest.mark.parametrize("arg_value", [True, False])
    def it_knows_whether_to_include_text_as_html_in_Table_metadata(
        self, arg_value: bool, opts_args: dict[str, Any]
    ):
        opts_args["infer_table_structure"] = arg_value
        opts = DocxPartitionerOptions(**opts_args)

        assert opts.infer_table_structure is arg_value

    # -- .increment_page_number() ----------------

    def it_generates_a_PageBreak_element_when_the_page_number_is_incremented(
        self, opts_args: dict[str, Any]
    ):
        opts = DocxPartitionerOptions(**opts_args)

        page_break_iter = opts.increment_page_number()

        assert isinstance(next(page_break_iter, None), PageBreak)
        assert opts.page_number == 2
        with pytest.raises(StopIteration):
            next(page_break_iter)

    def but_it_does_not_generate_a_PageBreak_element_when_include_page_breaks_option_is_off(
        self, opts_args: dict[str, Any]
    ):
        opts_args["include_page_breaks"] = False
        opts = DocxPartitionerOptions(**opts_args)

        page_break_iter = opts.increment_page_number()

        with pytest.raises(StopIteration):
            next(page_break_iter)
        assert opts.page_number == 2

    # -- .last_modified --------------------------

    def it_gets_last_modified_from_the_filesystem_when_file_path_is_provided(
        self, opts_args: dict[str, Any], get_last_modified_date_: Mock
    ):
        opts_args["file_path"] = "a/b/document.docx"
        get_last_modified_date_.return_value = "2024-04-02T20:32:35"
        opts = DocxPartitionerOptions(**opts_args)

        last_modified = opts.last_modified

        get_last_modified_date_.assert_called_once_with("a/b/document.docx")
        assert last_modified == "2024-04-02T20:32:35"

    def but_it_falls_back_to_None_for_the_last_modified_date_when_no_file_path_is_provided(
        self, opts_args: dict[str, Any]
    ):
        file = io.BytesIO(b"abcdefg")
        opts_args["file"] = file
        opts = DocxPartitionerOptions(**opts_args)

        assert opts.last_modified is None

    # -- .metadata_file_path ---------------------

    @pytest.mark.parametrize("file_path", ["u/v/w.docx", None])
    def it_uses_the_file_path_argument_when_provided(
        self, file_path: str | None, opts_args: dict[str, Any]
    ):
        opts_args["file_path"] = file_path
        opts = DocxPartitionerOptions(**opts_args)

        assert opts.metadata_file_path == file_path

    # -- ._metadata_page_number ------------------

    @pytest.mark.parametrize(
        ("page_count", "document_contains_pagebreaks", "expected_value"),
        [(7, True, 7), (1, False, None)],
    )
    def it_reports_None_when_no_rendered_page_breaks_are_found_in_document(
        self,
        request: FixtureRequest,
        opts_args: dict[str, Any],
        page_count: int,
        document_contains_pagebreaks: bool,
        expected_value: int | None,
    ):
        _document_contains_pagebreaks_prop_ = property_mock(
            request,
            DocxPartitionerOptions,
            "_document_contains_pagebreaks",
            return_value=document_contains_pagebreaks,
        )
        opts = DocxPartitionerOptions(**opts_args)
        opts._page_counter = page_count

        metadata_page_number = opts.metadata_page_number

        _document_contains_pagebreaks_prop_.assert_called_once_with()
        assert metadata_page_number is expected_value

    # -- .page_number ----------------------------

    def it_keeps_track_of_the_page_number(self, opts_args: dict[str, Any]):
        """In DOCX, page-number is the slide number."""
        opts = DocxPartitionerOptions(**opts_args)

        assert opts.page_number == 1
        list(opts.increment_page_number())
        assert opts.page_number == 2
        list(opts.increment_page_number())
        assert opts.page_number == 3

    def it_assigns_the_correct_page_number_when_starting_page_number_is_given(
        self, opts_args: dict[str, Any]
    ):
        opts = DocxPartitionerOptions(**opts_args, starting_page_number=3)

        assert opts.page_number == 3
        list(opts.increment_page_number())
        assert opts.page_number == 4

    # -- .strategy -------------------------------

    @pytest.mark.parametrize(
        ("arg_value", "expected_value"),
        [(None, "hi_res"), (PartitionStrategy.FAST, "fast"), (PartitionStrategy.HI_RES, "hi_res")],
    )
    def it_knows_which_partitioning_strategy_to_use(
        self, opts_args: dict[str, Any], arg_value: str, expected_value: str
    ):
        opts_args["strategy"] = arg_value
        opts = DocxPartitionerOptions(**opts_args)

        assert opts.strategy == expected_value

    # -- ._document_contains_pagebreaks ----------

    @pytest.mark.parametrize(
        ("file_name", "expected_value"), [("page-breaks.docx", True), ("teams_chat.docx", False)]
    )
    def it_knows_whether_the_document_contains_page_breaks(
        self, opts_args: dict[str, Any], file_name: str, expected_value: bool
    ):
        opts_args["file_path"] = example_doc_path(file_name)
        opts = DocxPartitionerOptions(**opts_args)

        assert opts._document_contains_pagebreaks is expected_value

    # -- ._docx_file -----------------------------

    def it_uses_the_path_to_open_the_presentation_when_file_path_is_provided(
        self, opts_args: dict[str, Any]
    ):
        opts_args["file_path"] = "l/m/n.docx"
        opts = DocxPartitionerOptions(**opts_args)

        assert opts._docx_file == "l/m/n.docx"

    def and_it_uses_a_BytesIO_file_to_replaces_a_SpooledTemporaryFile_provided(
        self, opts_args: dict[str, Any]
    ):
        with tempfile.SpooledTemporaryFile() as spooled_temp_file:
            spooled_temp_file.write(b"abcdefg")
            opts_args["file"] = spooled_temp_file
            opts = DocxPartitionerOptions(**opts_args)

            docx_file = opts._docx_file

            assert docx_file is not spooled_temp_file
            assert isinstance(docx_file, io.BytesIO)
            assert docx_file.getvalue() == b"abcdefg"

    def and_it_uses_the_provided_file_directly_when_not_a_SpooledTemporaryFile(
        self, opts_args: dict[str, Any]
    ):
        file = io.BytesIO(b"abcdefg")
        opts_args["file"] = file
        opts = DocxPartitionerOptions(**opts_args)

        docx_file = opts._docx_file

        assert docx_file is file
        assert isinstance(docx_file, io.BytesIO)
        assert docx_file.getvalue() == b"abcdefg"

    # -- ._validate() ----------------------------

    def it_raises_when_no_file_exists_at_file_path(self, opts_args: dict[str, Any]):
        opts_args["file_path"] = "l/m/n.docx"
        with pytest.raises(FileNotFoundError, match="no such file or directory: 'l/m/n.docx'"):
            DocxPartitionerOptions.load(**opts_args)

    def and_it_raises_when_the_file_at_file_path_is_not_a_ZIP_archive(
        self, opts_args: dict[str, Any]
    ):
        opts_args["file_path"] = example_doc_path("simple.doc")
        with pytest.raises(ValueError, match=r"not a ZIP archive \(so not a DOCX file\): "):
            DocxPartitionerOptions.load(**opts_args)

    def and_it_raises_when_the_file_like_object_is_not_a_ZIP_archive(
        self, opts_args: dict[str, Any]
    ):
        with open(example_doc_path("simple.doc"), "rb") as f:
            opts_args["file"] = f
            with pytest.raises(ValueError, match=r"not a ZIP archive \(so not a DOCX file\): "):
                DocxPartitionerOptions.load(**opts_args)

    def and_it_raises_when_neither_a_file_path_or_file_is_provided(self, opts_args: dict[str, Any]):
        with pytest.raises(ValueError, match="no DOCX document specified, either `filename` or "):
            DocxPartitionerOptions.load(**opts_args)

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture()
    def get_last_modified_date_(self, request: FixtureRequest) -> Mock:
        return function_mock(request, "unstructured.partition.docx.get_last_modified_date")


class Describe_DocxPartitioner:
    """Unit-test suite for `unstructured.partition.docx._DocxPartitioner`."""

    # -- table behaviors -------------------------------------------------------------------------

    def it_can_convert_a_table_to_html(self, opts_args: dict[str, Any]):
        opts = DocxPartitionerOptions(**opts_args)
        table = docx.Document(example_doc_path("docx-tables.docx")).tables[0]

        assert _DocxPartitioner(opts)._convert_table_to_html(table) == (
            "<table>"
            "<tr><td>Header Col 1</td><td>Header Col 2</td></tr>"
            "<tr><td>Lorem ipsum</td><td>A link example</td></tr>"
            "</table>"
        )

    def and_it_can_convert_a_nested_table_to_html(self, opts_args: dict[str, Any]):
        """
        Fixture table is:

            +---+-------------+---+
            | a |     >b<     | c |
            +---+-------------+---+
            |   | +-----+---+ |   |
            |   | |  e  | f | |   |
            | d | +-----+---+ | i |
            |   | | g&t | h | |   |
            |   | +-----+---+ |   |
            +---+-------------+---+
            | j |      k      | l |
            +---+-------------+---+
        """
        opts = DocxPartitionerOptions(**opts_args)
        table = docx.Document(example_doc_path("docx-tables.docx")).tables[1]

        # -- re.sub() strips out the extra padding inserted by tabulate --
        html = re.sub(r" +<", "<", _DocxPartitioner(opts)._convert_table_to_html(table))

        assert html == (
            "<table>"
            "<tr><td>a</td><td>&gt;b&lt;</td><td>c</td></tr>"
            "<tr><td>d</td><td>e f g&amp;t h</td><td>i</td></tr>"
            "<tr><td>j</td><td>k</td><td>l</td></tr>"
            "</table>"
        )

    def it_can_convert_a_table_to_plain_text(self, opts_args: dict[str, Any]):
        opts = DocxPartitionerOptions(**opts_args)
        table = docx.Document(example_doc_path("docx-tables.docx")).tables[0]

        assert " ".join(_DocxPartitioner(opts)._iter_table_texts(table)) == (
            "Header Col 1 Header Col 2 Lorem ipsum A link example"
        )

    def and_it_can_convert_a_nested_table_to_plain_text(self, opts_args: dict[str, Any]):
        """
        Fixture table is:

            +---+-------------+---+
            | a |     >b<     | c |
            +---+-------------+---+
            |   | +-----+---+ |   |
            |   | |  e  | f | |   |
            | d | +-----+---+ | i |
            |   | | g&t | h | |   |
            |   | +-----+---+ |   |
            +---+-------------+---+
            | j |      k      | l |
            +---+-------------+---+
        """
        opts = DocxPartitionerOptions(**opts_args)
        table = docx.Document(example_doc_path("docx-tables.docx")).tables[1]

        assert " ".join(_DocxPartitioner(opts)._iter_table_texts(table)) == (
            "a >b< c d e f g&t h i j k l"
        )

    def but_the_text_of_a_merged_cell_appears_only_once(self, opts_args: dict[str, Any]):
        """
        Fixture table is:

            +---+-------+
            | a | b     |
            |   +---+---+
            |   | c | d |
            +---+---+   |
            | e     |   |
            +-------+---+
        """
        opts = DocxPartitionerOptions(**opts_args)
        table = docx.Document(example_doc_path("docx-tables.docx")).tables[2]
        assert " ".join(_DocxPartitioner(opts)._iter_table_texts(table)) == "a b c d e"

    def it_can_partition_tables_with_incomplete_rows(self):
        """DOCX permits table rows to start late and end early.

        It is relatively rare in the wild, but DOCX tables are unique (as far as I know) in that
        they allow rows to start late, like in column 3, and end early, like the last cell is in
        column 5 of a 7 column table.

        A practical example might look like this:

                       +------+------+
                       | East | West |
            +----------+------+------+
            | Started  |  25  |  32  |
            +----------+------+------+
            | Finished |  17  |  21  |
            +----------+------+------+
        """
        elements = iter(partition_docx(example_doc_path("tables-with-incomplete-rows.docx")))

        e = next(elements)
        assert e.text.startswith("Example of DOCX table ")
        # --
        # ┌───┬───┐
        # │ a │ b │
        # ├───┼───┤
        # │ c │ d │
        # └───┴───┘
        e = next(elements)
        assert type(e).__name__ == "Table"
        assert e.text == "a b c d"
        assert e.metadata.text_as_html == (
            "<table><tr><td>a</td><td>b</td></tr><tr><td>c</td><td>d</td></tr></table>"
        )
        # --
        # ┌───┐
        # │ a │
        # ├───┼───┐
        # │ b │ c │
        # └───┴───┘
        e = next(elements)
        assert type(e).__name__ == "Table"
        assert e.text == "a b c", f"actual {e.text=}"
        assert e.metadata.text_as_html == (
            "<table><tr><td>a</td><td/></tr><tr><td>b</td><td>c</td></tr></table>"
        ), f"actual {e.metadata.text_as_html=}"
        # --
        # ┌───────┐
        # │   a   │
        # ├───┬───┼───┐
        # │ b │ c │ d │
        # └───┴───┴───┘
        e = next(elements)
        assert type(e).__name__ == "Table"
        assert e.text == "a b c d", f"actual {e.text=}"
        assert e.metadata.text_as_html == (
            "<table>"
            "<tr><td>a</td><td>a</td><td/></tr>"
            "<tr><td>b</td><td>c</td><td>d</td></tr>"
            "</table>"
        ), f"actual {e.metadata.text_as_html=}"
        # --
        # ┌───┬───┐
        # │   │ b │
        # │ a ├───┼───┐
        # │   │ c │ d │
        # └───┴───┴───┘
        e = next(elements)
        assert type(e).__name__ == "Table"
        assert e.text == "a b c d", f"actual {e.text=}"
        assert e.metadata.text_as_html == (
            "<table>"
            "<tr><td>a</td><td>b</td><td/></tr>"
            "<tr><td>a</td><td>c</td><td>d</td></tr>"
            "</table>"
        ), f"actual {e.metadata.text_as_html=}"
        # -- late-start, early-end, and >2 rows vertical span --
        # ┌───────┬───┬───┐
        # │   a   │ b │ c │
        # └───┬───┴───┼───┘
        #     │   d   │
        # ┌───┤       ├───┐
        # │ e │       │ f │
        # └───┤       ├───┘
        #     │       │
        #     └───────┘
        e = next(elements)
        assert type(e).__name__ == "Table"
        assert e.text == "a b c d e f", f"actual {e.text=}"
        assert e.metadata.text_as_html == (
            "<table>"
            "<tr><td>a</td><td>a</td><td>b</td><td>c</td></tr>"
            "<tr><td/><td>d</td><td>d</td><td/></tr>"
            "<tr><td>e</td><td>d</td><td>d</td><td>f</td></tr>"
            "<tr><td/><td>d</td><td>d</td><td/></tr>"
            "</table>"
        ), f"actual {e.metadata.text_as_html=}"
        # --
        # -- The table from the specimen file we received with the bug report. --
        e = next(elements)
        assert type(e).__name__ == "Table"
        assert e.text == "Data More Dato WTF? Strange Format", f"actual {e.text=}"
        assert e.metadata.text_as_html == (
            "<table>"
            "<tr><td>Data</td><td>Data</td><td/></tr>"
            "<tr><td>Data</td><td>Data</td><td/></tr>"
            "<tr><td>Data</td><td>Data</td><td/></tr>"
            "<tr><td/><td>More</td><td/></tr>"
            "<tr><td>Dato</td><td/></tr>"
            "<tr><td>WTF?</td><td>WTF?</td><td/></tr>"
            "<tr><td>Strange</td><td>Strange</td><td/></tr>"
            "<tr><td/><td>Format</td><td>Format</td></tr>"
            "</table>"
        ), f"actual {e.metadata.text_as_html=}"

    # -- page-break behaviors --------------------------------------------------------------------

    def it_places_page_breaks_precisely_where_they_occur(self, opts_args: dict[str, Any]):
        """Page-break behavior has some subtleties.

        * A hard page-break does not generate a PageBreak element (because that would double-count
          it). Word inserts a rendered page-break for the hard break at the effective location.
        * A (rendered) page-break mid-paragraph produces two elements, like `Text, PageBreak, Text`,
          so each Text (subclass) element gets the right page-number.
        * A rendered page-break mid-hyperlink produces two text elements, but the hyperlink itself
          is not split; the entire hyperlink goes on the page where the hyperlink starts, even
          though some of its text appears on the following page. The rest of the paragraph, after
          the hyperlink, appears on the following page.
        * Odd and even-page section starts can lead to two page-breaks, like an odd-page section
          start could go from page 3 to page 5 because 5 is the next odd page.
        """

        def str_repr(e: Element) -> str:
            """A more detailed `repr()` to aid debugging when assertion fails."""
            return f"{e.__class__.__name__}('{e}')"

        opts_args["file_path"] = example_doc_path("page-breaks.docx")
        opts = DocxPartitionerOptions(**opts_args)
        expected = [
            # -- page 1 --
            NarrativeText(
                "First page, tab here:\t"
                "followed by line-break here:\n"
                "here:\n"
                "and here:\n"
                "no-break hyphen here:-"
                "and hard page-break here>>"
            ),
            PageBreak(""),
            # -- page 2 --
            NarrativeText(
                "<<Text on second page. The font is big so it breaks onto third page--"
                "------------------here-->> <<but break falls inside link so text stays"
                " together."
            ),
            PageBreak(""),
            # -- page 3 --
            NarrativeText("Continuous section break here>>"),
            NarrativeText("<<followed by text on same page"),
            NarrativeText("Odd-page section break here>>"),
            PageBreak(""),
            # -- page 4 --
            PageBreak(""),
            # -- page 5 --
            NarrativeText("<<producing two page-breaks to get from page-3 to page-5."),
            NarrativeText(
                'Then text gets big again so a "natural" rendered page break happens again here>> '
            ),
            PageBreak(""),
            # -- page 6 --
            Text("<<and then more text proceeds."),
        ]

        elements = _DocxPartitioner.iter_document_elements(opts)

        for idx, e in enumerate(elements):
            assert e == expected[idx], (
                f"\n\nExpected: {str_repr(expected[idx])}"
                # --
                f"\n\nGot:      {str_repr(e)}\n"
            )

    # -- header/footer behaviors -----------------------------------------------------------------

    def it_includes_table_cell_text_in_Header_text(self, opts_args: dict[str, Any]):
        opts_args["file_path"] = example_doc_path("docx-hdrftr.docx")
        opts = DocxPartitionerOptions(**opts_args)
        partitioner = _DocxPartitioner(opts)
        section = partitioner._document.sections[0]

        header_iter = partitioner._iter_section_headers(section)

        element = next(header_iter)
        assert element.text == "First header para\nTable cell1 Table cell2\nLast header para"

    def it_includes_table_cell_text_in_Footer_text(self, opts_args: dict[str, Any]):
        """This case also verifies nested-table and merged-cell behaviors."""
        opts_args["file_path"] = example_doc_path("docx-hdrftr.docx")
        opts = DocxPartitionerOptions(**opts_args)
        partitioner = _DocxPartitioner(opts)
        section = partitioner._document.sections[0]

        footer_iter = partitioner._iter_section_footers(section)

        element = next(footer_iter)
        assert element.text == "para1\ncell1 a b c d e f\npara2"
