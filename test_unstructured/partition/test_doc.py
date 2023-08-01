import os

import docx
import pytest

from unstructured.documents.elements import (
    Address,
    ListItem,
    NarrativeText,
    Text,
    Title,
)
from unstructured.partition.common import convert_office_doc
from unstructured.partition.doc import partition_doc
from unstructured.partition.docx import partition_docx


@pytest.fixture()
def mock_document():
    document = docx.Document()

    document.add_paragraph("These are a few of my favorite things:", style="Heading 1")
    # NOTE(robinson) - this should get picked up as a list item due to the •
    document.add_paragraph("• Parrots", style="Normal")
    # NOTE(robinson) - this should get dropped because it's empty
    document.add_paragraph("• ", style="Normal")
    document.add_paragraph("Hockey", style="List Bullet")
    # NOTE(robinson) - this should get dropped because it's empty
    document.add_paragraph("", style="List Bullet")
    # NOTE(robinson) - this should get picked up as a title
    document.add_paragraph("Analysis", style="Normal")
    # NOTE(robinson) - this should get dropped because it is empty
    document.add_paragraph("", style="Normal")
    # NOTE(robinson) - this should get picked up as a narrative text
    document.add_paragraph("This is my first thought. This is my second thought.", style="Normal")
    document.add_paragraph("This is my third thought.", style="Body Text")
    # NOTE(robinson) - this should just be regular text
    document.add_paragraph("2023")
    # NOTE(robinson) - this should be an address
    document.add_paragraph("DOYLESTOWN, PA 18901")

    return document


@pytest.fixture()
def expected_elements():
    return [
        Title("These are a few of my favorite things:"),
        ListItem("Parrots"),
        ListItem("Hockey"),
        Title("Analysis"),
        NarrativeText("This is my first thought. This is my second thought."),
        NarrativeText("This is my third thought."),
        Text("2023"),
        Address("DOYLESTOWN, PA 18901"),
    ]


def test_partition_doc_from_filename(mock_document, expected_elements, tmpdir, capsys):
    docx_filename = os.path.join(tmpdir.dirname, "mock_document.docx")
    doc_filename = os.path.join(tmpdir.dirname, "mock_document.doc")
    mock_document.save(docx_filename)
    convert_office_doc(docx_filename, tmpdir.dirname, "doc")
    elements = partition_doc(filename=doc_filename)
    assert elements == expected_elements
    assert elements[0].metadata.filename == "mock_document.doc"
    assert elements[0].metadata.file_directory == tmpdir.dirname
    assert capsys.readouterr().out == ""
    assert capsys.readouterr().err == ""


def test_partition_doc_from_filename_with_metadata_filename(
    mock_document,
    expected_elements,
    tmpdir,
):
    docx_filename = os.path.join(tmpdir.dirname, "mock_document.docx")
    doc_filename = os.path.join(tmpdir.dirname, "mock_document.doc")
    mock_document.save(docx_filename)
    convert_office_doc(docx_filename, tmpdir.dirname, "doc")

    elements = partition_doc(filename=doc_filename, metadata_filename="test")
    assert elements == expected_elements
    assert all(element.metadata.filename == "test" for element in elements)


def test_partition_doc_matches_partition_docx(mock_document, expected_elements, tmpdir):
    docx_filename = os.path.join(tmpdir.dirname, "mock_document.docx")
    doc_filename = os.path.join(tmpdir.dirname, "mock_document.doc")
    mock_document.save(docx_filename)
    convert_office_doc(docx_filename, tmpdir.dirname, "doc")
    assert partition_doc(filename=doc_filename) == partition_docx(filename=docx_filename)


def test_partition_raises_with_missing_doc(mock_document, expected_elements, tmpdir):
    doc_filename = os.path.join(tmpdir.dirname, "asdf.doc")

    with pytest.raises(ValueError):
        partition_doc(filename=doc_filename)


def test_partition_doc_from_file_with_filter(mock_document, expected_elements, tmpdir, capsys):
    docx_filename = os.path.join(tmpdir.dirname, "mock_document.docx")
    doc_filename = os.path.join(tmpdir.dirname, "mock_document.doc")
    mock_document.save(docx_filename)
    convert_office_doc(docx_filename, tmpdir.dirname, "doc")

    with open(doc_filename, "rb") as f:
        elements = partition_doc(file=f, libre_office_filter="MS Word 2007 XML")
    assert elements == expected_elements
    assert capsys.readouterr().out == ""
    assert capsys.readouterr().err == ""
    for element in elements:
        assert element.metadata.filename is None


def test_partition_doc_from_file_with_no_filter(mock_document, expected_elements, tmpdir, capsys):
    docx_filename = os.path.join(tmpdir.dirname, "mock_document.docx")
    doc_filename = os.path.join(tmpdir.dirname, "mock_document.doc")
    mock_document.save(docx_filename)
    convert_office_doc(docx_filename, tmpdir.dirname, "doc")

    with open(doc_filename, "rb") as f:
        elements = partition_doc(file=f, libre_office_filter=None)
    assert elements == expected_elements
    assert capsys.readouterr().out == ""
    assert capsys.readouterr().err == ""
    for element in elements:
        assert element.metadata.filename is None


def test_partition_doc_from_file_with_metadata_filename(mock_document, tmpdir):
    docx_filename = os.path.join(tmpdir.dirname, "mock_document.docx")
    doc_filename = os.path.join(tmpdir.dirname, "mock_document.doc")
    mock_document.save(docx_filename)
    convert_office_doc(docx_filename, tmpdir.dirname, "doc")

    with open(doc_filename, "rb") as f:
        elements = partition_doc(file=f, metadata_filename="test")
    for element in elements:
        assert element.metadata.filename == "test"


def test_partition_doc_raises_with_both_specified(mock_document, tmpdir):
    docx_filename = os.path.join(tmpdir.dirname, "mock_document.docx")
    doc_filename = os.path.join(tmpdir.dirname, "mock_document.doc")
    mock_document.save(docx_filename)
    convert_office_doc(docx_filename, tmpdir.dirname, "doc")

    with open(doc_filename, "rb") as f, pytest.raises(ValueError):
        partition_doc(filename=doc_filename, file=f)


def test_partition_doc_raises_with_neither():
    with pytest.raises(ValueError):
        partition_doc()


def test_partition_doc_from_file_exclude_metadata(mock_document, tmpdir):
    docx_filename = os.path.join(tmpdir.dirname, "mock_document.docx")
    doc_filename = os.path.join(tmpdir.dirname, "mock_document.doc")
    mock_document.save(docx_filename)
    convert_office_doc(docx_filename, tmpdir.dirname, "doc")

    with open(doc_filename, "rb") as f:
        elements = partition_doc(file=f, include_metadata=False)

    assert elements[0].metadata.filetype is None
    assert elements[0].metadata.page_name is None
    assert elements[0].metadata.filename is None


def test_partition_doc_from_filename_exclude_metadata(mock_document, tmpdir):
    docx_filename = os.path.join(tmpdir.dirname, "mock_document.docx")
    doc_filename = os.path.join(tmpdir.dirname, "mock_document.doc")
    mock_document.save(docx_filename)
    convert_office_doc(docx_filename, tmpdir.dirname, "doc")

    elements = partition_doc(filename=doc_filename, include_metadata=False)

    assert elements[0].metadata.filetype is None
    assert elements[0].metadata.page_name is None
    assert elements[0].metadata.filename is None
