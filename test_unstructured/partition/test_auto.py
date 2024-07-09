# pyright: reportPrivateUsage=false

from __future__ import annotations

import json
import os
import pathlib
import tempfile
import warnings
from importlib import import_module
from typing import Callable, Iterator, cast
from unittest.mock import Mock, patch

import docx
import pytest
from docx.document import Document
from PIL import Image

from test_unstructured.partition.pdf_image.test_pdf import assert_element_extraction
from test_unstructured.partition.test_constants import (
    EXPECTED_TABLE,
    EXPECTED_TABLE_XLSX,
    EXPECTED_TEXT,
    EXPECTED_TEXT_XLSX,
    EXPECTED_TITLE,
)
from test_unstructured.unit_utils import (
    ANY,
    FixtureRequest,
    LogCaptureFixture,
    MonkeyPatch,
    example_doc_path,
    function_mock,
    method_mock,
)
from unstructured.chunking.title import chunk_by_title
from unstructured.cleaners.core import clean_extra_whitespace
from unstructured.documents.elements import (
    Address,
    Element,
    ElementMetadata,
    ListItem,
    NarrativeText,
    Table,
    TableChunk,
    Text,
    Title,
)
from unstructured.file_utils.filetype import FILETYPE_TO_MIMETYPE, FileType
from unstructured.partition import auto
from unstructured.partition.auto import _get_partition_with_extras, partition
from unstructured.partition.common import convert_office_doc
from unstructured.partition.utils.constants import PartitionStrategy
from unstructured.staging.base import elements_to_json

DIRECTORY = pathlib.Path(__file__).parent.resolve()
EXAMPLE_DOCS_DIRECTORY = os.path.join(DIRECTORY, "..", "..", "example-docs")

EXPECTED_EMAIL_OUTPUT = [
    NarrativeText(text="This is a test email to use for unit tests."),
    Title(text="Important points:"),
    ListItem(text="Roses are red"),
    ListItem(text="Violets are blue"),
]

EML_TEST_FILE = "eml/fake-email.eml"

is_in_docker = os.path.exists("/.dockerenv")


# ================================================================================================
# CSV
# ================================================================================================


@pytest.mark.skipif(is_in_docker, reason="Skipping this test in Docker container")
def test_auto_partition_csv_from_filename():
    elements = partition(example_doc_path("stanley-cups.csv"))

    assert clean_extra_whitespace(elements[0].text) == EXPECTED_TEXT
    assert elements[0].metadata.text_as_html == EXPECTED_TABLE
    assert elements[0].metadata.filetype == "text/csv"


@pytest.mark.skipif(is_in_docker, reason="Skipping this test in Docker container")
def test_auto_partition_csv_from_file():
    with open(example_doc_path("stanley-cups.csv"), "rb") as f:
        elements = partition(file=f)

    assert clean_extra_whitespace(elements[0].text) == EXPECTED_TEXT
    assert isinstance(elements[0], Table)
    assert elements[0].metadata.text_as_html == EXPECTED_TABLE
    assert elements[0].metadata.filetype == "text/csv"


# ================================================================================================
# DOC
# ================================================================================================


@pytest.mark.skipif(is_in_docker, reason="Skipping this test in Docker container")
@pytest.mark.parametrize(
    ("pass_metadata_filename", "content_type"),
    [(False, None), (False, "application/msword"), (True, "application/msword"), (True, None)],
)
def test_auto_partition_doc_with_filename(
    mock_docx_document: Document,
    expected_docx_elements: list[Element],
    tmp_path: pathlib.Path,
    pass_metadata_filename: bool,
    content_type: str | None,
):
    docx_file_path = str(tmp_path / "mock_document.docx")
    doc_file_path = str(tmp_path / "mock_document.doc")
    mock_docx_document.save(docx_file_path)
    convert_office_doc(docx_file_path, str(tmp_path), "doc")
    metadata_filename = doc_file_path if pass_metadata_filename else None
    elements = partition(
        filename=doc_file_path,
        metadata_filename=metadata_filename,
        content_type=content_type,
        strategy=PartitionStrategy.HI_RES,
    )
    assert elements == expected_docx_elements
    assert elements[0].metadata.filename == "mock_document.doc"
    assert elements[0].metadata.file_directory == str(tmp_path)


# NOTE(robinson) - the application/x-ole-storage mime type is not specific enough to
# determine that the file is an .doc document
@pytest.mark.xfail()
def test_auto_partition_doc_with_file(
    mock_docx_document: Document, expected_docx_elements: list[Element], tmp_path: pathlib.Path
):
    docx_filename = str(tmp_path / "mock_document.docx")
    doc_filename = str(tmp_path / "mock_document.doc")
    mock_docx_document.save(docx_filename)
    convert_office_doc(docx_filename, str(tmp_path), "doc")

    with open(doc_filename, "rb") as f:
        elements = partition(file=f, strategy=PartitionStrategy.HI_RES)
    assert elements == expected_docx_elements


# ================================================================================================
# DOCX
# ================================================================================================


@pytest.fixture()
def mock_docx_document():
    document = docx.Document()

    document.add_paragraph("These are a few of my favorite things:", style="Heading 1")
    # NOTE(robinson) - this should get picked up as a list item due to the •
    document.add_paragraph("• Parrots", style="Normal")
    document.add_paragraph("Hockey", style="List Bullet")
    # NOTE(robinson) - this should get picked up as a title
    document.add_paragraph("Analysis", style="Normal")
    # NOTE(robinson) - this should get dropped because it is empty
    document.add_paragraph("", style="Normal")
    # NOTE(robinson) - this should get picked up as a narrative text
    document.add_paragraph("This is my first thought. This is my second thought.", style="Normal")
    document.add_paragraph("This is my third thought.", style="Body Text")
    # NOTE(robinson) - this should just be regular text
    document.add_paragraph("2023")

    return document


@pytest.fixture()
def expected_docx_elements():
    return [
        Title("These are a few of my favorite things:"),
        ListItem("Parrots"),
        ListItem("Hockey"),
        Title("Analysis"),
        NarrativeText("This is my first thought. This is my second thought."),
        NarrativeText("This is my third thought."),
        Text("2023"),
    ]


def test_auto_partition_docx_with_filename(
    mock_docx_document: Document, expected_docx_elements: list[Element], tmp_path: pathlib.Path
):
    filename = str(tmp_path / "mock_document.docx")
    mock_docx_document.save(filename)

    elements = partition(filename=filename, strategy=PartitionStrategy.HI_RES)
    assert elements == expected_docx_elements
    assert elements[0].metadata.filename == os.path.basename(filename)


def test_auto_partition_docx_with_file(
    mock_docx_document: Document, expected_docx_elements: list[Element], tmp_path: pathlib.Path
):
    filename = str(tmp_path / "mock_document.docx")
    mock_docx_document.save(filename)

    with open(filename, "rb") as f:
        elements = partition(file=f, strategy=PartitionStrategy.HI_RES)
    assert elements == expected_docx_elements


@pytest.mark.parametrize("file_name", ["simple.docx", "simple.doc", "simple.odt"])
@pytest.mark.parametrize(
    "strategy",
    [
        PartitionStrategy.AUTO,
        PartitionStrategy.FAST,
        PartitionStrategy.HI_RES,
        PartitionStrategy.OCR_ONLY,
    ],
)
def test_partition_forwards_strategy_arg_to_partition_docx_and_its_brokers(
    request: FixtureRequest, file_name: str, strategy: str
):
    """The `strategy` arg value received by `partition()` is received by `partition_docx().

    To do this in the brokering-partitioner cases (DOC, ODT) it must make its way to
    `partition_doc()` or `partition_odt()` which must then forward it to `partition_docx()`. This
    test makes sure it made it all the way.

    Note this is 3 file-types X 4 strategies = 12 test-cases.
    """
    from unstructured.partition.docx import _DocxPartitioner

    def fake_iter_document_elements(self: _DocxPartitioner) -> Iterator[Element]:
        yield Text(f"strategy=={self._opts.strategy}")

    _iter_elements_ = method_mock(
        request,
        _DocxPartitioner,
        "_iter_document_elements",
        side_effect=fake_iter_document_elements,
    )

    (element,) = partition(example_doc_path(file_name), strategy=strategy)

    _iter_elements_.assert_called_once_with(ANY)
    assert element.text == f"strategy=={strategy}"


# ================================================================================================
# EML
# ================================================================================================


def test_auto_partition_email_from_filename():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, EML_TEST_FILE)
    elements = partition(filename=filename, strategy=PartitionStrategy.HI_RES)
    assert len(elements) > 0
    assert elements == EXPECTED_EMAIL_OUTPUT
    assert elements[0].metadata.filename == os.path.basename(filename)
    assert elements[0].metadata.file_directory == os.path.split(filename)[0]


def test_auto_partition_email_from_file():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, EML_TEST_FILE)
    with open(filename, "rb") as f:
        elements = partition(file=f, strategy=PartitionStrategy.HI_RES)
    assert len(elements) > 0
    assert elements == EXPECTED_EMAIL_OUTPUT


def test_auto_partition_email_from_file_rb():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, EML_TEST_FILE)
    with open(filename, "rb") as f:
        elements = partition(file=f, strategy=PartitionStrategy.HI_RES)
    assert len(elements) > 0
    assert elements == EXPECTED_EMAIL_OUTPUT


def test_auto_partition_eml_add_signature_to_metadata():
    elements = partition(filename="example-docs/eml/signed-doc.p7s")
    assert len(elements) == 1
    assert elements[0].text == "This is a test"
    assert elements[0].metadata.signature == "<SIGNATURE>\n"


# ================================================================================================
# EPUB
# ================================================================================================


def test_auto_partition_epub_from_filename():
    filename = os.path.join(DIRECTORY, "..", "..", "example-docs", "winter-sports.epub")
    elements = partition(filename=filename, strategy=PartitionStrategy.HI_RES)
    assert len(elements) > 0
    assert elements[0].text.startswith("The Project Gutenberg eBook of Winter Sports")


def test_auto_partition_epub_from_file():
    filename = os.path.join(DIRECTORY, "..", "..", "example-docs", "winter-sports.epub")
    with open(filename, "rb") as f:
        elements = partition(file=f, strategy=PartitionStrategy.HI_RES)
    assert len(elements) > 0
    assert elements[0].text.startswith("The Project Gutenberg eBook of Winter Sports")


# ================================================================================================
# HTML
# ================================================================================================


@pytest.mark.parametrize(
    ("pass_metadata_filename", "content_type"),
    [(False, None), (False, "text/html"), (True, "text/html"), (True, None)],
)
def test_auto_partition_html_from_filename(pass_metadata_filename: bool, content_type: str | None):
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "example-10k.html")
    metadata_filename = filename if pass_metadata_filename else None
    elements = partition(
        filename=filename,
        metadata_filename=metadata_filename,
        content_type=content_type,
        strategy=PartitionStrategy.HI_RES,
    )
    assert len(elements) > 0
    assert elements[0].metadata.filename == os.path.basename(filename)
    assert elements[0].metadata.file_directory == os.path.split(filename)[0]


@pytest.mark.parametrize(
    ("pass_metadata_filename", "content_type"),
    [(False, None), (False, "text/html"), (True, "text/html"), (True, None)],
)
def test_auto_partition_html_from_file(pass_metadata_filename: bool, content_type: str | None):
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake-html.html")
    metadata_filename = filename if pass_metadata_filename else None
    with open(filename, "rb") as f:
        elements = partition(
            file=f,
            metadata_filename=metadata_filename,
            content_type=content_type,
            strategy=PartitionStrategy.HI_RES,
        )
    assert len(elements) > 0


def test_auto_partition_html_from_file_rb():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake-html.html")
    with open(filename, "rb") as f:
        elements = partition(file=f, strategy=PartitionStrategy.HI_RES)
    assert len(elements) > 0


def test_auto_partition_html_pre_from_file():
    elements = partition(example_doc_path("fake-html-pre.htm"))

    assert len(elements) > 0
    assert "PageBreak" not in [elem.category for elem in elements]
    assert clean_extra_whitespace(elements[0].text).startswith("[107th Congress Public Law 56]")
    assert isinstance(elements[0], NarrativeText)
    assert elements[0].metadata.filetype == "text/html"
    assert elements[0].metadata.filename == "fake-html-pre.htm"


# ================================================================================================
# IMAGE
# ================================================================================================


@pytest.mark.parametrize(
    ("pass_metadata_filename", "content_type"),
    [(False, None), (False, "image/jpeg"), (True, "image/jpeg"), (True, None)],
)
def test_auto_partition_image(pass_metadata_filename: bool, content_type: str | None):
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "layout-parser-paper-fast.jpg")
    metadata_filename = filename if pass_metadata_filename else None
    elements = partition(
        filename=filename,
        metadata_filename=metadata_filename,
        content_type=content_type,
        strategy=PartitionStrategy.AUTO,
    )

    # should be same result as test_partition_image_default_strategy_hi_res() in test_image.py
    title = "LayoutParser: A Unified Toolkit for Deep Learning Based Document Image Analysis"
    idx = 2
    assert elements[idx].text == title
    assert elements[idx].metadata.coordinates is not None


@pytest.mark.parametrize("extract_image_block_to_payload", [False, True])
def test_auto_partition_image_element_extraction(extract_image_block_to_payload: bool):
    extract_image_block_types = ["Image", "Table"]

    with tempfile.TemporaryDirectory() as tmpdir:
        elements = partition(
            filename=example_doc_path("embedded-images-tables.jpg"),
            extract_image_block_types=extract_image_block_types,
            extract_image_block_to_payload=extract_image_block_to_payload,
            extract_image_block_output_dir=tmpdir,
        )

        assert_element_extraction(
            elements, extract_image_block_types, extract_image_block_to_payload, tmpdir
        )


@pytest.mark.parametrize(
    ("pass_metadata_filename", "content_type"),
    [(False, None), (False, "image/jpeg"), (True, "image/jpeg"), (True, None)],
)
def test_auto_partition_jpg(pass_metadata_filename: bool, content_type: str | None):
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "layout-parser-paper-fast.jpg")
    metadata_filename = filename if pass_metadata_filename else None
    elements = partition(
        filename=filename,
        metadata_filename=metadata_filename,
        content_type=content_type,
        strategy=PartitionStrategy.AUTO,
    )
    assert len(elements) > 0


@pytest.mark.parametrize(
    ("pass_metadata_filename", "content_type"),
    [(False, None), (False, "image/jpeg"), (True, "image/jpeg"), (True, None)],
)
def test_auto_partition_jpg_from_file(pass_metadata_filename: bool, content_type: str | None):
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "layout-parser-paper-fast.jpg")
    metadata_filename = filename if pass_metadata_filename else None
    with open(filename, "rb") as f:
        elements = partition(
            file=f,
            metadata_filename=metadata_filename,
            content_type=content_type,
            strategy=PartitionStrategy.AUTO,
        )
    assert len(elements) > 0


def test_partition_image_with_bmp_with_auto(tmp_path: pathlib.Path):
    bmp_filename = str(tmp_path / "example.bmp")
    with Image.open(example_doc_path("layout-parser-paper-with-table.jpg")) as img:
        img.save(bmp_filename)

    elements = partition(
        filename=bmp_filename,
        strategy=PartitionStrategy.HI_RES,
    )

    table = [e.metadata.text_as_html for e in elements if e.metadata.text_as_html]
    assert len(table) == 1
    assert "<table><thead><tr>" in table[0]
    assert "</thead><tbody><tr>" in table[0]


# ================================================================================================
# JSON
# ================================================================================================


# NOTE(robinson) - skipping this test with docker image to avoid putting the
# test fixtures into the image
@pytest.mark.skipif(is_in_docker, reason="Skipping this test in Docker container")
def test_auto_partitioned_json_output_maintains_consistency_with_fixture_elements():
    """Test auto-processing an unstructured json output file by filename."""
    original_file_name = "spring-weather.html"
    json_file_path = (
        pathlib.Path(DIRECTORY).parents[1]
        / "test_unstructured_ingest"
        / "expected-structured-output"
        / "azure"
        / f"{original_file_name}.json"
    )
    with open(json_file_path) as json_f:
        expected_result = json.load(json_f)

    partitioning_result = json.loads(
        cast(
            str,
            elements_to_json(
                partition(
                    filename=str(json_file_path),
                    # -- use the original file name to get the same element IDs (hashes) --
                    metadata_filename=original_file_name,
                    strategy=PartitionStrategy.HI_RES,
                )
            ),
        )
    )
    for elem in partitioning_result:
        elem.pop("metadata")
    for elem in expected_result:
        elem.pop("metadata")
    assert expected_result == partitioning_result


def test_auto_partition_json_raises_with_unprocessable_json(tmp_path: pathlib.Path):
    # NOTE(robinson) - This is unprocessable because it is not a list of dicts,
    # per the Unstructured ISD format
    text = '{"hi": "there"}'

    filename = str(tmp_path / "unprocessable.json")
    with open(filename, "w") as f:
        f.write(text)

    with pytest.raises(ValueError):
        partition(filename=filename)


@pytest.mark.xfail(
    reason="parsed as text not json, https://github.com/Unstructured-IO/unstructured/issues/492",
)
def test_auto_partition_json_from_file():
    """Test auto-processing an unstructured json output file by file handle."""
    filename = os.path.join(
        EXAMPLE_DOCS_DIRECTORY,
        "..",
        "test_unstructured_ingest",
        "expected-structured-output",
        "azure-blob-storage",
        "spring-weather.html.json",
    )
    with open(filename) as json_f:
        json_data = json.load(json_f)
    with open(filename, "rb") as partition_f:
        json_elems = json.loads(
            cast(
                str,
                elements_to_json(partition(file=partition_f, strategy=PartitionStrategy.HI_RES)),
            )
        )
    for elem in json_elems:
        # coordinates are always in the element data structures, even if None
        elem.pop("coordinates")
        elem.pop("coordinate_system")
    assert json_data == json_elems


def test_auto_partition_works_with_unstructured_jsons():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "spring-weather.html.json")
    elements = partition(filename=filename, strategy=PartitionStrategy.HI_RES)
    assert elements[0].text == "News Around NOAA"


def test_auto_partition_works_with_unstructured_jsons_from_file():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "spring-weather.html.json")
    with open(filename, "rb") as f:
        elements = partition(file=f, strategy=PartitionStrategy.HI_RES)
    assert elements[0].text == "News Around NOAA"


# ================================================================================================
# MD
# ================================================================================================


def test_partition_md_works_with_embedded_html():
    url = "https://raw.githubusercontent.com/Unstructured-IO/unstructured/main/README.md"
    elements = partition(url=url, content_type="text/markdown", strategy=PartitionStrategy.HI_RES)
    assert "unstructured" in elements[0].text


# ================================================================================================
# MSG
# ================================================================================================


EXPECTED_MSG_OUTPUT = [
    NarrativeText(text="This is a test email to use for unit tests."),
    Title(text="Important points:"),
    ListItem(text="Roses are red"),
    ListItem(text="Violets are blue"),
]


def test_auto_partition_msg_from_filename():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake-email.msg")
    elements = partition(filename=filename, strategy=PartitionStrategy.HI_RES)
    assert elements == EXPECTED_MSG_OUTPUT


# ================================================================================================
# ODT
# ================================================================================================


def test_auto_partition_odt_from_filename():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake.odt")
    elements = partition(filename=filename, strategy=PartitionStrategy.HI_RES)
    assert elements[0] == Title("Lorem ipsum dolor sit amet.")


def test_auto_partition_odt_from_file():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake.odt")
    with open(filename, "rb") as f:
        elements = partition(file=f, strategy=PartitionStrategy.HI_RES)

    assert elements[0] == Title("Lorem ipsum dolor sit amet.")


# ================================================================================================
# ORG
# ================================================================================================


def test_auto_partition_org_from_filename():
    elements = partition(example_doc_path("README.org"))

    assert elements[0] == Title("Example Docs")
    assert elements[0].metadata.filetype == "text/org"


def test_auto_partition_org_from_file():
    with open(example_doc_path("README.org"), "rb") as f:
        elements = partition(file=f, content_type="text/org")

    assert elements[0] == Title("Example Docs")
    assert elements[0].metadata.filetype == "text/org"


# ================================================================================================
# PDF
# ================================================================================================


@pytest.mark.parametrize(
    ("pass_metadata_filename", "content_type"),
    [(False, None), (False, "application/pdf"), (True, "application/pdf"), (True, None)],
)
def test_auto_partition_pdf_from_filename(
    request: FixtureRequest, pass_metadata_filename: bool, content_type: str | None
):
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "layout-parser-paper-fast.pdf")
    metadata_filename = filename if pass_metadata_filename else None

    elements = partition(
        filename=filename,
        metadata_filename=metadata_filename,
        content_type=content_type,
        strategy=PartitionStrategy.HI_RES,
    )

    # NOTE(alan): Xfail since new model skips the word Zejiang
    request.applymarker(pytest.mark.xfail)

    idx = 3
    assert isinstance(elements[idx], Title)
    assert elements[idx].text.startswith("LayoutParser")

    assert elements[idx].metadata.filename == os.path.basename(filename)
    assert elements[idx].metadata.file_directory == os.path.split(filename)[0]

    idx += 1
    assert isinstance(elements[idx], NarrativeText)
    assert elements[idx].text.startswith("Zejiang Shen")


def test_auto_partition_pdf_uses_table_extraction():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "layout-parser-paper-fast.pdf")
    with patch(
        "unstructured.partition.pdf_image.ocr.process_file_with_ocr",
    ) as mock_process_file_with_model:
        partition(filename, pdf_infer_table_structure=True, strategy=PartitionStrategy.HI_RES)
        assert mock_process_file_with_model.call_args[1]["infer_table_structure"]


def test_auto_partition_pdf_with_fast_strategy(monkeypatch: MonkeyPatch):
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "layout-parser-paper-fast.pdf")

    mock_return = [NarrativeText("Hello there!")]
    with patch.object(auto, "partition_pdf", return_value=mock_return) as mock_partition:
        mock_partition_with_extras_map = {"pdf": mock_partition}
        monkeypatch.setattr(auto, "PARTITION_WITH_EXTRAS_MAP", mock_partition_with_extras_map)
        partition(filename=filename, strategy=PartitionStrategy.FAST)

    mock_partition.assert_called_once_with(
        filename=filename,
        file=None,
        url=None,
        strategy=PartitionStrategy.FAST,
        languages=None,
        metadata_filename=None,
        include_page_breaks=False,
        infer_table_structure=False,
        extract_images_in_pdf=False,
        extract_image_block_types=None,
        extract_image_block_output_dir=None,
        extract_image_block_to_payload=False,
        hi_res_model_name=None,
        date_from_file_object=False,
        starting_page_number=1,
    )


@pytest.mark.parametrize(
    ("pass_metadata_filename", "content_type"),
    [(False, None), (False, "application/pdf"), (True, "application/pdf"), (True, None)],
)
def test_auto_partition_pdf_from_file(
    request: FixtureRequest, pass_metadata_filename: bool, content_type: str | None
):
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "layout-parser-paper-fast.pdf")
    metadata_filename = filename if pass_metadata_filename else None

    with open(filename, "rb") as f:
        elements = partition(
            file=f,
            metadata_filename=metadata_filename,
            content_type=content_type,
            strategy=PartitionStrategy.HI_RES,
        )

    # NOTE(alan): Xfail since new model skips the word Zejiang
    request.applymarker(pytest.mark.xfail)

    idx = 3
    assert isinstance(elements[idx], Title)
    assert elements[idx].text.startswith("LayoutParser")

    idx += 1
    assert isinstance(elements[idx], NarrativeText)
    assert elements[idx].text.startswith("Zejiang Shen")


def test_partition_pdf_does_not_raise_warning():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "layout-parser-paper-fast.pdf")
    # NOTE(robinson): This is the recommended way to check that no warning is emitted,
    # per the pytest docs.
    # ref: https://docs.pytest.org/en/7.0.x/how-to/capture-warnings.html
    #      #additional-use-cases-of-warnings-in-tests
    with warnings.catch_warnings():
        warnings.simplefilter("error")
        partition(filename=filename, strategy=PartitionStrategy.HI_RES)


@pytest.mark.parametrize("extract_image_block_to_payload", [False, True])
def test_auto_partition_pdf_element_extraction(extract_image_block_to_payload: bool):
    extract_image_block_types = ["Image", "Table"]

    with tempfile.TemporaryDirectory() as tmpdir:
        elements = partition(
            example_doc_path("embedded-images-tables.pdf"),
            extract_image_block_types=extract_image_block_types,
            extract_image_block_to_payload=extract_image_block_to_payload,
            extract_image_block_output_dir=tmpdir,
        )

        assert_element_extraction(
            elements, extract_image_block_types, extract_image_block_to_payload, tmpdir
        )


# ================================================================================================
# PPT
# ================================================================================================


@pytest.mark.skipif(is_in_docker, reason="Skipping this test in Docker container")
def test_auto_partition_ppt_from_filename():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake-power-point.ppt")
    elements = partition(filename=filename, strategy=PartitionStrategy.HI_RES)
    assert elements == EXPECTED_PPTX_OUTPUT
    assert elements[0].metadata.filename == os.path.basename(filename)
    assert elements[0].metadata.file_directory == os.path.split(filename)[0]


# ================================================================================================
# PPTX
# ================================================================================================


EXPECTED_PPTX_OUTPUT = [
    Title(text="Adding a Bullet Slide"),
    ListItem(text="Find the bullet slide layout"),
    ListItem(text="Use _TextFrame.text for first bullet"),
    ListItem(text="Use _TextFrame.add_paragraph() for subsequent bullets"),
    NarrativeText(text="Here is a lot of text!"),
    NarrativeText(text="Here is some text in a text box!"),
]


def test_auto_partition_pptx_from_filename():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake-power-point.pptx")
    elements = partition(filename=filename, strategy=PartitionStrategy.HI_RES)
    assert elements == EXPECTED_PPTX_OUTPUT
    assert elements[0].metadata.filename == os.path.basename(filename)
    assert elements[0].metadata.file_directory == os.path.split(filename)[0]


@pytest.mark.parametrize("file_name", ["simple.pptx", "fake-power-point.ppt"])
@pytest.mark.parametrize(
    "strategy",
    [
        PartitionStrategy.AUTO,
        PartitionStrategy.FAST,
        PartitionStrategy.HI_RES,
        PartitionStrategy.OCR_ONLY,
    ],
)
def test_partition_forwards_strategy_arg_to_partition_pptx_and_its_brokers(
    request: FixtureRequest, file_name: str, strategy: str
):
    """The `strategy` arg value received by `partition()` is received by `partition_pptx().

    To do this in the brokering-partitioner case (PPT) the strategy argument must make its way to
    `partition_ppt()` which must then forward it to `partition_pptx()`. This test makes sure it
    made it all the way.

    Note this is 2 file-types X 4 strategies = 8 test-cases.
    """
    from unstructured.partition.pptx import _PptxPartitioner

    def fake_iter_presentation_elements(self: _PptxPartitioner) -> Iterator[Element]:
        yield Text(f"strategy=={self._opts.strategy}")

    _iter_elements_ = method_mock(
        request,
        _PptxPartitioner,
        "_iter_presentation_elements",
        side_effect=fake_iter_presentation_elements,
    )

    (element,) = partition(example_doc_path(file_name), strategy=strategy)

    _iter_elements_.assert_called_once_with(ANY)
    assert element.text == f"strategy=={strategy}"


# ================================================================================================
# RST
# ================================================================================================


def test_auto_partition_rst_from_filename():
    elements = partition(example_doc_path("README.rst"))

    assert elements[0] == Title("Example Docs")
    assert elements[0].metadata.filetype == "text/x-rst"


def test_auto_partition_rst_from_file():
    with open(example_doc_path("README.rst"), "rb") as f:
        elements = partition(file=f, content_type="text/x-rst")

    assert elements[0] == Title("Example Docs")
    assert elements[0].metadata.filetype == "text/x-rst"


# ================================================================================================
# RTF
# ================================================================================================


def test_auto_partition_rtf_from_filename():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake-doc.rtf")
    elements = partition(filename=filename, strategy=PartitionStrategy.HI_RES)
    assert elements[0] == Title("My First Heading")


# ================================================================================================
# TSV
# ================================================================================================


@pytest.mark.skipif(is_in_docker, reason="Skipping this test in Docker container")
def test_auto_partition_tsv_from_filename():
    elements = partition(example_doc_path("stanley-cups.tsv"))

    assert clean_extra_whitespace(elements[0].text) == EXPECTED_TEXT
    assert elements[0].metadata.text_as_html == EXPECTED_TABLE
    assert elements[0].metadata.filetype == "text/tsv"


# ================================================================================================
# TXT
# ================================================================================================


EXPECTED_TEXT_OUTPUT = [
    NarrativeText(text="This is a test document to use for unit tests."),
    Address(text="Doylestown, PA 18901"),
    Title(text="Important points:"),
    ListItem(text="Hamburgers are delicious"),
    ListItem(text="Dogs are the best"),
    ListItem(text="I love fuzzy blankets"),
]


def test_auto_partition_text_from_filename():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake-text.txt")
    elements = partition(filename=filename, strategy=PartitionStrategy.HI_RES)
    assert len(elements) > 0
    assert elements == EXPECTED_TEXT_OUTPUT
    assert elements[0].metadata.filename == os.path.basename(filename)
    assert elements[0].metadata.file_directory == os.path.split(filename)[0]


def test_auto_partition_text_from_file():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake-text.txt")
    with open(filename, "rb") as f:
        elements = partition(file=f, strategy=PartitionStrategy.HI_RES)
    assert len(elements) > 0
    assert elements == EXPECTED_TEXT_OUTPUT


# ================================================================================================
# XLS
# ================================================================================================


EXPECTED_XLS_TEXT_LEN = 550


EXPECTED_XLS_INITIAL_45_CLEAN_TEXT = "MC What is 2+2? 4 correct 3 incorrect MA What"

EXPECTED_XLS_TABLE = (
    """<table border="1" class="dataframe">
  <tbody>
    <tr>
      <td>MC</td>
      <td>What is 2+2?</td>
      <td>4</td>
      <td>correct</td>
      <td>3</td>
      <td>incorrect</td>
      <td></td>
      <td></td>
      <td></td>
    </tr>
    <tr>
      <td>MA</td>
      <td>What C datatypes are 8 bits? (assume i386)</td>
      <td>int</td>
      <td></td>
      <td>float</td>
      <td></td>
      <td>double</td>
      <td></td>
      <td>char</td>
    </tr>
    <tr>
      <td>TF</td>
      <td>Bagpipes are awesome.</td>
      <td>true</td>
      <td></td>
      <td></td>
      <td></td>
      <td></td>
      <td></td>
      <td></td>
    </tr>
    <tr>
      <td>ESS</td>
      <td>How have the original Henry Hornbostel buildings """
    """influenced campus architecture and design in the last 30 years?</td>
      <td></td>
      <td></td>
      <td></td>
      <td></td>
      <td></td>
      <td></td>
      <td></td>
    </tr>
    <tr>
      <td>ORD</td>
      <td>Rank the following in their order of operation.</td>
      <td>Parentheses</td>
      <td>Exponents</td>
      <td>Division</td>
      <td>Addition</td>
      <td></td>
      <td></td>
      <td></td>
    </tr>
    <tr>
      <td>FIB</td>
      <td>The student activities fee is</td>
      <td>95</td>
      <td>dollars for students enrolled in</td>
      <td>19</td>
      <td>units or more,</td>
      <td></td>
      <td></td>
      <td></td>
    </tr>
    <tr>
      <td>MAT</td>
      <td>Match the lower-case greek letter with its capital form.</td>
      <td>λ</td>
      <td>Λ</td>
      <td>α</td>
      <td>γ</td>
      <td>Γ</td>
      <td>φ</td>
      <td>Φ</td>
    </tr>
  </tbody>
</table>"""
)


@pytest.mark.skipif(is_in_docker, reason="Skipping this test in Docker container")
def test_auto_partition_xls_from_filename():
    elements = partition(
        example_doc_path("tests-example.xls"), include_header=False, skip_infer_table_types=[]
    )

    assert sum(isinstance(element, Table) for element in elements) == 2
    assert len(elements) == 14

    assert clean_extra_whitespace(elements[0].text)[:45] == EXPECTED_XLS_INITIAL_45_CLEAN_TEXT
    # NOTE(crag): if the beautifulsoup4 package is installed, some (but not all) additional
    # whitespace is removed, so the expected text length is less than is the case
    # when beautifulsoup4 is *not* installed. E.g.
    # "\n\n\nMA\nWhat C datatypes are 8 bits" vs.
    # '\n  \n    \n      MA\n      What C datatypes are 8 bits?... "
    assert len(elements[0].text) == EXPECTED_XLS_TEXT_LEN
    assert elements[0].metadata.text_as_html == EXPECTED_XLS_TABLE


# ================================================================================================
# XLSX
# ================================================================================================


EXPECTED_XLSX_FILETYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def test_auto_partition_xlsx_from_filename():
    elements = partition(
        example_doc_path("stanley-cups.xlsx"), include_header=False, skip_infer_table_types=[]
    )

    assert sum(isinstance(element, Table) for element in elements) == 2
    assert sum(isinstance(element, Title) for element in elements) == 2
    assert len(elements) == 4

    assert clean_extra_whitespace(elements[0].text) == EXPECTED_TITLE
    assert clean_extra_whitespace(elements[1].text) == EXPECTED_TEXT_XLSX
    assert elements[1].metadata.text_as_html == EXPECTED_TABLE_XLSX
    assert elements[1].metadata.page_number == 1
    assert elements[1].metadata.filetype == EXPECTED_XLSX_FILETYPE


def test_auto_partition_xlsx_from_file():
    with open(example_doc_path("stanley-cups.xlsx"), "rb") as f:
        elements = partition(file=f, include_header=False, skip_infer_table_types=[])

    assert sum(isinstance(element, Table) for element in elements) == 2
    assert sum(isinstance(element, Title) for element in elements) == 2
    assert len(elements) == 4

    assert clean_extra_whitespace(elements[0].text) == EXPECTED_TITLE
    assert clean_extra_whitespace(elements[1].text) == EXPECTED_TEXT_XLSX
    assert elements[1].metadata.text_as_html == EXPECTED_TABLE_XLSX
    assert elements[1].metadata.page_number == 1
    assert elements[1].metadata.filetype == EXPECTED_XLSX_FILETYPE


def test_auto_partition_respects_starting_page_number_argument_for_xlsx():
    elements = partition("example-docs/stanley-cups.xlsx", starting_page_number=3)
    assert elements[1].metadata.page_number == 3


# ================================================================================================
# XML
# ================================================================================================


def test_auto_partition_xml_from_filename():
    file_path = example_doc_path("factbook.xml")

    elements = partition(file_path, xml_keep_tags=False, metadata_filename=file_path)

    assert elements[0].text == "United States"
    assert elements[0].metadata.filename == "factbook.xml"


def test_auto_partition_xml_from_file():
    with open(example_doc_path("factbook.xml"), "rb") as f:
        elements = partition(file=f, xml_keep_tags=False)

    assert elements[0].text == "United States"


def test_auto_partition_xml_from_filename_with_tags():
    elements = partition(example_doc_path("factbook.xml"), xml_keep_tags=True)

    assert "<leader>Joe Biden</leader>" in elements[0].text
    assert elements[0].metadata.filename == "factbook.xml"


def test_auto_partition_xml_from_file_with_tags():
    with open(example_doc_path("factbook.xml"), "rb") as f:
        elements = partition(file=f, xml_keep_tags=True)

    assert "<leader>Joe Biden</leader>" in elements[0].text


# ================================================================================================
# FILE_TYPE NOT RECOGNIZED OR NOT SUPPORTED
# ================================================================================================


def test_auto_partition_raises_with_bad_type(request: FixtureRequest):
    detect_filetype_ = function_mock(
        request, "unstructured.partition.auto.detect_filetype", return_value=None
    )

    with pytest.raises(ValueError):
        partition(filename="made-up.fake", strategy=PartitionStrategy.HI_RES)

    detect_filetype_.assert_called_once_with(
        content_type=None, encoding=None, file=None, file_filename=None, filename="made-up.fake"
    )


# ================================================================================================
# LOAD FROM URL
# ================================================================================================


def test_auto_partition_from_url():
    url = "https://raw.githubusercontent.com/Unstructured-IO/unstructured/main/LICENSE.md"
    elements = partition(url=url, content_type="text/plain", strategy=PartitionStrategy.HI_RES)
    assert elements[0] == Title("Apache License")
    assert elements[0].metadata.url == url


def test_auto_partition_from_url_with_rfc9110_content_type():
    url = "https://raw.githubusercontent.com/Unstructured-IO/unstructured/main/LICENSE.md"
    elements = partition(
        url=url, content_type="text/plain; charset=utf-8", strategy=PartitionStrategy.HI_RES
    )
    assert elements[0] == Title("Apache License")
    assert elements[0].metadata.url == url


def test_auto_partition_from_url_without_providing_content_type():
    url = "https://raw.githubusercontent.com/Unstructured-IO/unstructured/main/LICENSE.md"
    elements = partition(url=url, strategy=PartitionStrategy.HI_RES)
    assert elements[0] == Title("Apache License")
    assert elements[0].metadata.url == url


def test_auto_partition_warns_if_header_set_and_not_url(caplog: LogCaptureFixture):
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, EML_TEST_FILE)
    partition(
        filename=filename, headers={"Accept": "application/pdf"}, strategy=PartitionStrategy.HI_RES
    )
    assert caplog.records[0].levelname == "WARNING"


def test_partition_timeout_gets_routed():
    class CallException(Exception):
        pass

    mock_ocr_func = Mock(side_effect=CallException("Function called!"))
    with patch("unstructured.partition.auto.file_and_type_from_url", mock_ocr_func), pytest.raises(
        CallException
    ):
        auto.partition(url="fake_url", request_timeout=326)
    kwargs = mock_ocr_func.call_args.kwargs
    assert "request_timeout" in kwargs
    assert kwargs["request_timeout"] == 326


# ================================================================================================
# OTHER ARGS
# ================================================================================================

# -- chunking_strategy ----------------------------------------------------


def test_add_chunking_strategy_on_partition_auto():
    filename = "example-docs/example-10k-1p.html"
    elements = partition(filename)
    chunk_elements = partition(filename, chunking_strategy="by_title")
    chunks = chunk_by_title(elements)
    assert chunk_elements != elements
    assert chunk_elements == chunks


def test_add_chunking_strategy_on_partition_auto_respects_max_chars():
    filename = "example-docs/example-10k-1p.html"

    # default chunk size in chars is 200
    partitioned_table_elements_200_chars = [
        e
        for e in partition(
            filename,
            chunking_strategy="by_title",
            max_characters=200,
            combine_text_under_n_chars=5,
        )
        if isinstance(e, (Table, TableChunk))
    ]

    partitioned_table_elements_5_chars = [
        e
        for e in partition(
            filename,
            chunking_strategy="by_title",
            max_characters=5,
            combine_text_under_n_chars=5,
        )
        if isinstance(e, (Table, TableChunk))
    ]

    elements = partition(filename)

    table_elements = [e for e in elements if isinstance(e, Table)]

    assert len(partitioned_table_elements_5_chars) != len(table_elements)
    assert len(partitioned_table_elements_200_chars) != len(table_elements)

    # trailing whitespace is stripped from the first chunk, leaving only a checkbox character
    assert len(partitioned_table_elements_5_chars[0].text) == 1
    # but the second chunk is the full 5 characters
    assert len(partitioned_table_elements_5_chars[1].text) == 5
    assert len(cast(str, partitioned_table_elements_5_chars[0].metadata.text_as_html)) == 5

    # the first table element is under 200 chars so doesn't get chunked!
    assert table_elements[0] == partitioned_table_elements_200_chars[0]
    assert len(partitioned_table_elements_200_chars[0].text) < 200
    assert len(partitioned_table_elements_200_chars[1].text) == 198
    assert len(cast(str, partitioned_table_elements_200_chars[1].metadata.text_as_html)) == 200


def test_add_chunking_strategy_chars_on_partition_auto_adds_is_continuation():
    filename = "example-docs/example-10k-1p.html"

    table_elements = [e for e in partition(filename) if isinstance(e, Table)]
    table_chunks = [
        e
        for e in partition(filename, chunking_strategy="by_title")
        if isinstance(e, (Table, TableChunk))
    ]

    assert table_elements != table_chunks

    i = 0
    for chunk in table_chunks:
        # have to reset the counter to 0 here when we encounter a Table element
        if not isinstance(chunk, TableChunk):
            i = 0
        if i > 0 and isinstance(chunk, TableChunk):
            assert chunk.metadata.is_continuation is True
            i += 1


# -- detect_language_per_element ------------------------------------------


def test_partition_respects_detect_language_per_element_arg():
    filename = "example-docs/language-docs/eng_spa_mult.txt"
    elements = partition(filename=filename, detect_language_per_element=True)
    langs = [element.metadata.languages for element in elements]
    assert langs == [["eng"], ["spa", "eng"], ["eng"], ["eng"], ["spa"]]


# -- languages ------------------------------------------------------------


@pytest.mark.parametrize(
    "file_extension",
    [
        "doc",
        "docx",
        "eml",
        "epub",
        "html",
        "md",
        "odt",
        "org",
        "ppt",
        "pptx",
        "rst",
        "rtf",
        "txt",
        "xml",
    ],
)
def test_partition_respects_language_arg(file_extension: str):
    elements = partition(
        example_doc_path(f"language-docs/eng_spa_mult.{file_extension}"), languages=["deu"]
    )
    assert all(element.metadata.languages == ["deu"] for element in elements)


# -- include_page_breaks --------------------------------------------------


def test_auto_with_page_breaks():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "layout-parser-paper-fast.pdf")
    elements = partition(
        filename=filename, include_page_breaks=True, strategy=PartitionStrategy.HI_RES
    )
    assert "PageBreak" in [elem.category for elem in elements]


# -- metadata_filename ----------------------------------------------------


def test_auto_partition_metadata_filename():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake-text.txt")
    with open(filename, "rb") as f:
        elements = partition(file=f, metadata_filename=filename)
    assert elements[0].metadata.filename == os.path.split(filename)[-1]


def test_auto_partition_warns_about_file_filename_deprecation(caplog: LogCaptureFixture):
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake-text.txt")
    with open(filename, "rb") as f:
        elements = partition(file=f, file_filename=filename)
    assert elements[0].metadata.filename == os.path.split(filename)[-1]
    assert "WARNING" in caplog.text
    assert "The file_filename kwarg will be deprecated" in caplog.text


def test_auto_partition_raises_with_file_and_metadata_filename():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "fake-text.txt")
    with open(filename, "rb") as f, pytest.raises(ValueError):
        partition(file=f, file_filename=filename, metadata_filename=filename)


# -- ocr_languages --------------------------------------------------------


def test_auto_partition_formats_languages_for_tesseract():
    filename = "example-docs/chi_sim_image.jpeg"
    with patch(
        "unstructured.partition.pdf_image.ocr.process_file_with_ocr",
    ) as mock_process_file_with_ocr:
        partition(filename, strategy=PartitionStrategy.HI_RES, languages=["zh"])
        _, kwargs = mock_process_file_with_ocr.call_args_list[0]
        assert "ocr_languages" in kwargs
        assert kwargs["ocr_languages"] == "chi_sim+chi_sim_vert+chi_tra+chi_tra_vert"


@pytest.mark.parametrize(("languages", "ocr_languages"), [(["auto"], ""), (["eng"], "")])
def test_auto_partition_ignores_empty_string_for_ocr_languages(
    languages: list[str], ocr_languages: str
):
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "book-war-and-peace-1p.txt")
    elements = partition(
        filename=filename,
        strategy=PartitionStrategy.OCR_ONLY,
        ocr_languages=ocr_languages,
        languages=languages,
    )
    assert elements[0].metadata.languages == ["eng"]


def test_auto_partition_warns_with_ocr_languages(caplog: LogCaptureFixture):
    filename = "example-docs/chevron-page.pdf"
    partition(filename=filename, strategy=PartitionStrategy.HI_RES, ocr_languages="eng")
    assert "The ocr_languages kwarg will be deprecated" in caplog.text


# -- skip_infer_table_types -----------------------------------------------


@pytest.mark.parametrize(
    ("skip_infer_table_types", "filename", "has_text_as_html_field"),
    [
        (["xlsx"], "stanley-cups.xlsx", False),
        ([], "stanley-cups.xlsx", True),
        (["odt"], "fake.odt", False),
        ([], "fake.odt", True),
    ],
)
def test_auto_partition_respects_skip_infer_table_types(
    skip_infer_table_types: list[str], filename: str, has_text_as_html_field: bool
):
    with open(example_doc_path(filename), "rb") as f:
        table_elements = [
            e
            for e in partition(file=f, skip_infer_table_types=skip_infer_table_types)
            if isinstance(e, Table)
        ]
        for table_element in table_elements:
            table_element_has_text_as_html_field = (
                hasattr(table_element.metadata, "text_as_html")
                and table_element.metadata.text_as_html is not None
            )
            assert table_element_has_text_as_html_field == has_text_as_html_field


# ================================================================================================
# METADATA BEHAVIORS
# ================================================================================================

# -- .filetype ------------------------------------------------------------

supported_filetypes = [t for t in FileType if t not in (FileType.UNK, FileType.ZIP, FileType.XLS)]

FILETYPE_TO_MODULE = {
    FileType.JPG: "image",
    FileType.PNG: "image",
    FileType.HEIC: "image",
    FileType.TXT: "text",
    FileType.EML: "email",
}


@pytest.mark.parametrize(
    ("content_type", "routing_func", "expected"),
    [
        ("text/csv", "csv", "text/csv"),
        ("text/html", "html", "text/html"),
        ("jdsfjdfsjkds", "pdf", None),
    ],
)
def test_auto_adds_filetype_to_metadata(
    request: FixtureRequest,
    content_type: str,
    routing_func: str,
    expected: str | None,
    monkeypatch: MonkeyPatch,
):
    partition_fn_ = function_mock(
        request,
        f"unstructured.partition.auto.partition_{routing_func}",
        return_value=[Text("text 1"), Text("text 2")],
    )
    mock_partition_with_extras_map = {routing_func: partition_fn_}
    monkeypatch.setattr(auto, "PARTITION_WITH_EXTRAS_MAP", mock_partition_with_extras_map)

    elements = partition("example-docs/layout-parser-paper-fast.pdf", content_type=content_type)

    assert len(elements) == 2
    assert all(el.metadata.filetype == expected for el in elements)


@pytest.mark.parametrize(
    ("content_type", "expected"),
    [
        ("application/pdf", FILETYPE_TO_MIMETYPE[FileType.PDF]),
        (None, FILETYPE_TO_MIMETYPE[FileType.PDF]),
    ],
)
def test_auto_filetype_overrides_file_specific(
    request: FixtureRequest, content_type: str | None, expected: str, monkeypatch: MonkeyPatch
):
    pdf_metadata = ElementMetadata(filetype="imapdf")
    partition_pdf_ = function_mock(
        request,
        "unstructured.partition.auto.partition_pdf",
        return_value=[Text("text 1", metadata=pdf_metadata), Text("text 2", metadata=pdf_metadata)],
    )
    mock_partition_with_extras_map = {"pdf": partition_pdf_}
    monkeypatch.setattr(auto, "PARTITION_WITH_EXTRAS_MAP", mock_partition_with_extras_map)

    elements = partition("example-docs/layout-parser-paper-fast.pdf", content_type=content_type)

    assert len(elements) == 2
    assert all(el.metadata.filetype == expected for el in elements)


@pytest.mark.parametrize("filetype", supported_filetypes)
def test_file_specific_produces_correct_filetype(filetype: FileType):
    if filetype in auto.IMAGE_FILETYPES or filetype in (FileType.WAV, FileType.EMPTY):
        pytest.skip()
    extension = filetype.name.lower()
    filetype_module = FILETYPE_TO_MODULE.get(filetype, extension)
    fun_name = "partition_" + filetype_module
    module = import_module(f"unstructured.partition.{filetype_module}")
    fun = getattr(module, fun_name)
    for file in pathlib.Path("example-docs").iterdir():
        if file.is_file() and file.suffix == f".{extension}":
            elements = fun(str(file))
            assert all(
                el.metadata.filetype == FILETYPE_TO_MIMETYPE[filetype]
                for el in elements
                if el.metadata.filetype is not None
            )
            break


# -- .languages -----------------------------------------------------------


def test_auto_partition_element_metadata_user_provided_languages():
    filename = "example-docs/chevron-page.pdf"
    elements = partition(filename=filename, strategy=PartitionStrategy.OCR_ONLY, languages=["eng"])
    assert elements[0].metadata.languages == ["eng"]


def test_partition_languages_incorrectly_defaults_to_English(tmp_path: pathlib.Path):
    # -- We don't totally rely on langdetect for short text, so text like the following that is
    # -- in German will be labeled as English.
    german = "Ein kurzer Satz."
    filepath = str(tmp_path / "short-german.txt")
    with open(filepath, "w") as f:
        f.write(german)
    elements = partition(filepath)
    assert elements[0].metadata.languages == ["eng"]


def test_partition_languages_default_to_None():
    filename = "example-docs/handbook-1p.docx"
    elements = partition(filename=filename, detect_language_per_element=True)
    # PageBreak and other elements with no text will have `None` for `languages`
    none_langs = [element for element in elements if element.metadata.languages is None]
    assert none_langs[0].text == ""


def test_partition_default_does_not_overwrite_other_defaults():
    """`partition()` ["eng"] default does not overwrite ["auto"] default in other partitioners."""
    # the default for `languages` is ["auto"] in partiton_text
    from unstructured.partition.text import partition_text

    # Use a document that is primarily in a language other than English
    filename = "example-docs/language-docs/UDHR_first_article_all.txt"
    text_elements = partition_text(filename)
    assert text_elements[0].metadata.languages != ["eng"]

    auto_elements = partition(filename)
    assert auto_elements[0].metadata.languages != ["eng"]
    assert auto_elements[0].metadata.languages == text_elements[0].metadata.languages


# ================================================================================================
# MISCELLANEOUS BEHAVIORS
# ================================================================================================


def test_auto_partition_works_on_empty_filename():
    assert partition(example_doc_path("empty.txt")) == []


def test_auto_partition_works_on_empty_file():
    with open(example_doc_path("empty.txt"), "rb") as f:
        assert partition(file=f) == []


def test_get_partition_with_extras_prompts_for_install_if_missing():
    partition_with_extras_map: dict[str, Callable[..., list[Element]]] = {}
    with pytest.raises(ImportError) as exception_info:
        _get_partition_with_extras("pdf", partition_with_extras_map)

    msg = str(exception_info.value)
    assert 'Install the pdf dependencies with pip install "unstructured[pdf]"' in msg
